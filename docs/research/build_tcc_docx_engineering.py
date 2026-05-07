from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
RESEARCH_DIR = ROOT / "docs" / "research"
INPUT_DOCX = Path(r"C:\Users\PEDRO\Downloads\tcc_heal_redisus_abnt_2026 (1).docx")
OUTPUT_DOCX = RESEARCH_DIR / "tcc_heal_redisus_abnt_2026_engenharia_editavel.docx"
OUTPUT_DOWNLOADS = Path(r"C:\Users\PEDRO\Downloads\tcc_heal_redisus_abnt_2026_engenharia_editavel.docx")
FIG_DIR = RESEARCH_DIR / "figures" / "software"
FONT = "Arial"


doc = Document(INPUT_DOCX)


def idx_of(target):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph._element is target._element:
            return index
    raise ValueError("Parágrafo não encontrado pelo elemento OOXML.")


def format_runs(paragraph, size=12, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = FONT
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic


def set_text(paragraph, text, size=12, bold=False, align=None):
    try:
        paragraph.style = doc.styles["Normal"]
    except KeyError:
        pass
    paragraph.text = text
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = False
    format_runs(paragraph, size=size, bold=bold)
    return paragraph


def format_body(paragraph, justify=True):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Inches(0.49)
    paragraph.paragraph_format.space_after = Pt(6)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    format_runs(paragraph, 12)


def format_caption(paragraph):
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_runs(paragraph, 12)


def format_source(paragraph):
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_runs(paragraph, 10)


def format_heading(paragraph, level=1):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_runs(paragraph, 12, bold=(level <= 2))


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def remove_section_break(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return
    sect_pr = p_pr.find(qn("w:sectPr"))
    if sect_pr is not None:
        p_pr.remove(sect_pr)


def find_para(predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph.text.strip()):
            return paragraph
    raise RuntimeError("Parágrafo não encontrado.")


def replace_between_elements(start, end, entries, title_text=None):
    if title_text is not None:
        set_text(start, title_text, 12, True, WD_ALIGN_PARAGRAPH.CENTER)
        remove_section_break(start)
    paragraphs = list(doc.paragraphs)
    start_index = idx_of(start)
    end_index = idx_of(end)
    for paragraph in paragraphs[start_index + 1 : end_index]:
        delete_paragraph(paragraph)
    for text, indent in entries:
        paragraph = end.insert_paragraph_before(text)
        try:
            paragraph.style = doc.styles["Normal"]
        except KeyError:
            pass
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.left_indent = Inches(0.25 if indent else 0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = False
        format_runs(paragraph, 12)


def parse_toc(path):
    entries = []
    for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        match = re.search(
            r"\\contentsline \{chapter\}\{\\chapternumberline \{([^}]+)\}(.+?)\}\{(\d+)\}",
            line,
        )
        if match:
            number, title, page = match.groups()
            entries.append((f"{number} {title}\t{page}", 0))
            continue
        match = re.search(
            r"\\contentsline \{section\}\{\\numberline \{([^}]+)\}(.+?)\}\{(\d+)\}",
            line,
        )
        if match:
            number, title, page = match.groups()
            entries.append((f"{number} {title}\t{page}", 1))
    return entries


def parse_list(path, prefix):
    entries = []
    for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        match = re.search(
            r"\\numberline \{([^}]+)\}\{(?:\\ignorespaces )?(.+?)\}\}\{(\d+)\}",
            line,
        )
        if match:
            number, title, page = match.groups()
            entries.append((f"{prefix} {number} – {title}\t{page}", 0))
    return entries


def add_para(anchor, text, kind="body"):
    paragraph = anchor.insert_paragraph_before(text)
    if kind == "h1":
        paragraph.paragraph_format.page_break_before = True
        format_heading(paragraph, 1)
    elif kind == "h2":
        format_heading(paragraph, 2)
    elif kind == "caption":
        format_caption(paragraph)
    elif kind == "source":
        format_source(paragraph)
    else:
        format_body(paragraph)
    return paragraph


def add_table(anchor, caption, rows):
    add_para(anchor, caption, "caption")
    table = doc.add_table(rows=1, cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = doc.styles["Normal Table"]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for column_index, value in enumerate(rows[0]):
        table.rows[0].cells[column_index].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for column_index, value in enumerate(row):
            cells[column_index].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.first_line_indent = None
                format_runs(paragraph, 9.5)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            format_runs(paragraph, 9.5, bold=True)
    anchor._p.addprevious(table._tbl)
    add_para(anchor, "Fonte: Elaboração própria.", "source")


def add_figure(anchor, filename, caption, source, width=6.0):
    paragraph = anchor.insert_paragraph_before("")
    paragraph.paragraph_format.first_line_indent = None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(FIG_DIR / filename), width=Inches(width))
    add_para(anchor, caption, "caption")
    add_para(anchor, source, "source")


# Atualiza resumo e palavras-chave.
summary_pt = next(
    (paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("Este trabalho apresenta a análise")),
    None,
)
if summary_pt:
    summary_pt.text = summary_pt.text.replace(
        "matriz de rastreabilidade, plano de testes",
        "matriz de rastreabilidade, modelagem de Engenharia de Software com casos de uso e diagramas UML, plano de testes",
    )
    format_body(summary_pt)

kw_pt = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("Palavras-chave:")), None)
if kw_pt:
    kw_pt.text = (
        "Palavras-chave: HEAL+; REDISUS; feridas crônicas; requisitos; saúde digital; "
        "engenharia de software; UML; inteligência artificial; visão computacional; LGPD; ética em pesquisa."
    )
    format_body(kw_pt, justify=False)

summary_en = next(
    (paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("This undergraduate thesis presents")),
    None,
)
if summary_en:
    summary_en.text = summary_en.text.replace(
        "a traceability matrix, a plan for public-image experiments",
        "a traceability matrix, Software Engineering modeling with use cases and UML diagrams, a plan for public-image experiments",
    )
    format_body(summary_en)

kw_en = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("Keywords:")), None)
if kw_en:
    kw_en.text = (
        "Keywords: HEAL+; REDISUS; chronic wounds; requirements; digital health; software engineering; "
        "UML; artificial intelligence; computer vision; data protection; research ethics."
    )
    format_body(kw_en, justify=False)


# Atualiza lista de siglas do DOCX-base.
for table in doc.tables:
    table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if "Technology Readiness Level" in table_text and "Unified Modeling Language" not in table_text:
        row = table.add_row()
        row.cells[0].text = "UML"
        row.cells[1].text = "Unified Modeling Language"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(2)
                format_runs(paragraph, 12)
        break


# Substitui listas estáticas incompletas por listas completas geradas pelo LaTeX.
toc_entries = parse_toc(RESEARCH_DIR / "tcc_heal_redisus_abnt_2026.toc")
lof_entries = parse_list(RESEARCH_DIR / "tcc_heal_redisus_abnt_2026.lof", "Figura")
lot_entries = parse_list(RESEARCH_DIR / "tcc_heal_redisus_abnt_2026.lot", "Tabela")

fig_first = find_para(lambda text: text.startswith("Figura 1"))
fig_title = list(doc.paragraphs)[idx_of(fig_first) - 1]
list_tables = find_para(lambda text: text == "LISTA DE TABELAS")
replace_between_elements(fig_title, list_tables, lof_entries, title_text="LISTA DE FIGURAS")

list_tables = find_para(lambda text: text == "LISTA DE TABELAS")
summary_title = find_para(lambda text: text == "SUMÁRIO")
replace_between_elements(list_tables, summary_title, lot_entries, title_text="LISTA DE TABELAS")

summary_title = find_para(lambda text: text == "SUMÁRIO")
intro_heading = find_para(lambda text: text.startswith("1 INTRODUÇÃO"))
replace_between_elements(summary_title, intro_heading, toc_entries, title_text="SUMÁRIO")

for title in ("LISTA DE FIGURAS", "LISTA DE TABELAS", "SUMÁRIO"):
    paragraph = find_para(lambda text, title=title: text == title)
    paragraph.paragraph_format.page_break_before = True

intro_heading = find_para(lambda text: text.startswith("1 INTRODUÇÃO"))
intro_heading.paragraph_format.page_break_before = True


# Insere Engenharia de Software antes dos testes.
anchor = find_para(lambda text: text.startswith("16 TESTES COM BANCOS DE IMAGENS PÚBLICOS"))

add_para(anchor, "16 ENGENHARIA DE SOFTWARE E MODELAGEM DA SOLUÇÃO", "h1")
add_para(
    anchor,
    "A modelagem de software do HEAL+ / REDISUS foi organizada para complementar o levantamento de requisitos "
    "técnicos, clínicos e éticos. Enquanto os requisitos definem o que a solução deve realizar, os modelos de "
    "Engenharia de Software descrevem, de forma visual e rastreável, os atores, fluxos, classes, entidades e "
    "interações que sustentam a proposta. Essa representação reduz ambiguidades, auxilia a comunicação entre "
    "equipe técnica e área clínica e favorece a verificação de coerência entre escopo, funcionalidades e regras "
    "de negócio (FERNANDES; MACHADO, 2017; PRESSMAN; MAXIM, 2016; GUEDES, 2018; WAZLAWICK, 2013).",
)
add_para(
    anchor,
    "A inclusão desta seção não altera o escopo assistencial do trabalho. O HEAL+ permanece definido como módulo "
    "de apoio ao registro, ao acompanhamento longitudinal e à análise evolutiva de feridas, com recursos de "
    "inteligência artificial tratados como assistivos, experimentais e dependentes de validação. Portanto, os "
    "diagramas apresentados devem ser lidos como documentação conceitual e acadêmica da solução, não como "
    "evidência de validação clínica ou de disponibilidade integral em produção.",
)

add_para(anchor, "16.1 Visão geral da modelagem", "h2")
add_para(
    anchor,
    "A modelagem foi organizada em cinco perspectivas: funcional, estrutural, comportamental, de interação e de "
    "dados. A perspectiva funcional é representada pelos atores e casos de uso; a estrutural, pelo diagrama de "
    "classes; a comportamental, pelos diagramas de atividades; a de interação, pelos diagramas de sequência; e a "
    "de dados, pelo modelo entidade-relacionamento conceitual. Essa organização permite relacionar os requisitos "
    "do TCC com a experiência esperada de uso do sistema e com os componentes conceituais necessários ao registro clínico.",
)
add_para(
    anchor,
    "O fluxo principal considerado para a modelagem envolve autenticação, seleção ou cadastro de paciente, registro "
    "da ferida, inclusão de imagem clínica, marcação de ROI, execução opcional de análise assistiva por IA, registro "
    "de avaliação clínica, armazenamento de evidências, geração de relatório e acompanhamento longitudinal. Quando "
    "houver uso de bases públicas ou imagens reais de pacientes, a modelagem também prevê separação entre uso "
    "assistencial, uso experimental e uso de pesquisa, respeitando LGPD, aprovação ética, TCLE e anonimização quando aplicável.",
)

add_para(anchor, "16.2 Atores do sistema", "h2")
add_table(
    anchor,
    "Tabela 9 – Atores do sistema HEAL+ / REDISUS",
    [
        ["Ator", "Descrição", "Tipo"],
        ["Profissional de saúde", "Usuário principal responsável por registrar pacientes, feridas, imagens, avaliações e relatórios.", "Principal"],
        ["Enfermeiro ou estomaterapeuta", "Profissional com atuação direta no acompanhamento de feridas, avaliação longitudinal e registro clínico estruturado.", "Principal"],
        ["Médico ou dermatologista", "Consulta histórico, imagens, relatórios e registros evolutivos para apoio à avaliação clínica.", "Principal"],
        ["Administrador", "Gerencia usuários, perfis, permissões, parâmetros e auditoria.", "Principal"],
        ["Pesquisador", "Acessa bases públicas, dados anonimizados, métricas e experimentos documentados.", "Principal"],
        ["Paciente", "Titular dos dados e sujeito do cuidado; no escopo atual, aparece como ator indireto ou futuro.", "Indireto"],
        ["Sistema de IA", "Processa imagem, ROI e parâmetros experimentais, retornando resultado assistivo e limitações.", "Secundário"],
        ["Sistema FHIR futuro", "Representa possibilidade de interoperabilidade clínica em versão posterior.", "Externo"],
        ["Banco de dados e armazenamento", "Mantém pacientes, feridas, imagens, avaliações, relatórios, resultados e logs.", "Técnico"],
    ],
)
add_para(
    anchor,
    "O paciente é tratado como sujeito central do cuidado e titular dos dados, mas não como usuário assistencial pleno "
    "no escopo atual. Caso um portal do paciente seja desenvolvido em versão futura, seus casos de uso deverão ser "
    "especificados separadamente, com regras próprias de consentimento, autenticação, linguagem de aviso e governança de dados.",
)

add_para(anchor, "16.3 Casos de uso principais", "h2")
add_table(
    anchor,
    "Tabela 10 – Casos de uso do HEAL+ / REDISUS",
    [
        ["Código", "Caso de uso", "Ator principal", "Status no escopo atual"],
        ["UC01", "Autenticar usuário", "Profissional, administrador ou pesquisador", "Parcialmente implementado"],
        ["UC02", "Gerenciar usuários e permissões", "Administrador", "Parcialmente implementado"],
        ["UC03", "Gerenciar paciente", "Profissional de saúde", "Parcialmente implementado"],
        ["UC04", "Registrar lesão ou ferida", "Profissional de saúde", "Parcialmente implementado"],
        ["UC05", "Enviar ou capturar imagem clínica", "Profissional de saúde", "Parcialmente implementado"],
        ["UC06", "Marcar ROI na imagem", "Profissional de saúde", "Experimental"],
        ["UC07", "Executar análise assistiva por IA", "Profissional e sistema de IA", "Experimental"],
        ["UC08", "Registrar avaliação clínica da ferida", "Profissional de saúde", "Parcialmente implementado"],
        ["UC09", "Consultar histórico evolutivo", "Profissional de saúde", "Parcialmente implementado"],
        ["UC10", "Comparar evolução por imagens", "Profissional de saúde", "Planejado/parcial"],
        ["UC11", "Gerar relatório de acompanhamento", "Profissional de saúde", "Parcialmente implementado"],
        ["UC12", "Registrar plano de cuidado ou conduta", "Profissional de saúde", "Planejado/parcial"],
        ["UC13", "Registrar follow-up", "Profissional de saúde", "Planejado/parcial"],
        ["UC14", "Registrar experimento com base pública", "Pesquisador", "Experimental"],
        ["UC15", "Aplicar pré-processamento com OpenCV", "Pesquisador ou sistema de IA", "Experimental"],
        ["UC16", "Exportar dados ou relatório", "Profissional ou pesquisador", "Planejado/parcial"],
        ["UC17", "Exportar recurso FHIR", "Sistema", "Planejado"],
        ["UC18", "Anonimizar dados para pesquisa", "Pesquisador ou administrador", "Planejado"],
        ["UC19", "Consultar logs e auditoria", "Administrador", "Planejado/parcial"],
    ],
)
add_table(
    anchor,
    "Tabela 11 – Documentação resumida dos casos de uso centrais",
    [
        ["UC", "Objetivo", "Fluxo resumido", "Requisitos"],
        ["UC01", "Controlar acesso ao sistema", "Usuário informa credenciais; a API valida autenticação; o perfil é verificado; o acesso é liberado conforme permissões e o evento é registrado.", "RF02, RF14, RF19"],
        ["UC03", "Manter cadastro mínimo do paciente", "Profissional seleciona novo paciente ou paciente existente; o sistema valida dados mínimos e vincula o registro ao contexto assistencial.", "RF03, RNF02, RNF03"],
        ["UC04", "Registrar ferida ou lesão", "Profissional informa tipo, localização, evolução, observações e vínculo com o paciente.", "RF04, RF05, RC03"],
        ["UC05", "Vincular imagem clínica", "Profissional associa imagem ao atendimento ou à ferida, preservando rastreabilidade, data e vínculo clínico.", "RF05, RNF03"],
        ["UC06", "Marcar ROI", "Profissional delimita região de interesse para análise visual, relatório ou experimento.", "RF06, RF07, RN06"],
        ["UC07", "Executar IA assistiva", "Sistema processa imagem ou ROI, registra modelo, versão, parâmetros e limitações, e apresenta resultado para revisão humana.", "RF15, RF16, RN15"],
        ["UC11", "Gerar relatório", "Sistema consolida dados clínicos, imagens, observações, resultados assistivos e aviso de limitação clínica.", "RF10, RC15, RN11"],
        ["UC18", "Anonimizar dados para pesquisa", "Usuário autorizado prepara dados para uso científico, removendo identificadores diretos e preservando rastreabilidade ética.", "RF17, RF19, RN09"],
    ],
)

add_para(anchor, "16.4 Diagramas da solução", "h2")
add_para(
    anchor,
    "As figuras a seguir apresentam a modelagem visual da solução. Para manter o documento tecnicamente correto, os "
    "diagramas indicam intenção de projeto, responsabilidades e fluxos esperados, sem afirmar que todos os recursos "
    "estejam validados clinicamente ou prontos para uso assistencial.",
)
add_figure(anchor, "use_case_diagram.png", "Figura 2 – Diagrama geral de casos de uso do HEAL+ / REDISUS", "Fonte: Elaboração própria, com base na modelagem funcional do projeto.", 6.2)
add_figure(anchor, "class_diagram.png", "Figura 3 – Diagrama de classes conceitual do HEAL+ / REDISUS", "Fonte: Elaboração própria, com base na modelagem estrutural da solução.", 6.2)
add_figure(anchor, "activity_login.png", "Figura 4 – Diagrama de atividades do fluxo de login", "Fonte: Elaboração própria.", 4.5)
add_figure(anchor, "activity_assessment.png", "Figura 5 – Diagrama de atividades do registro de avaliação da ferida", "Fonte: Elaboração própria.", 4.6)
add_figure(anchor, "activity_roi_ai_report.png", "Figura 6 – Diagrama de atividades do fluxo de ROI, IA assistiva e relatório", "Fonte: Elaboração própria.", 4.6)
add_figure(anchor, "sequence_login.png", "Figura 7 – Diagrama de sequência do fluxo de login", "Fonte: Elaboração própria.", 6.2)
add_figure(anchor, "sequence_patient_wound.png", "Figura 8 – Diagrama de sequência do gerenciamento de paciente e ferida", "Fonte: Elaboração própria.", 6.2)
add_figure(anchor, "sequence_roi_ai_report.png", "Figura 9 – Diagrama de sequência do fluxo de ROI, IA assistiva e relatório", "Fonte: Elaboração própria.", 6.2)
add_figure(anchor, "conceptual_er.png", "Figura 10 – Modelo entidade-relacionamento conceitual do HEAL+ / REDISUS", "Fonte: Elaboração própria, com base na modelagem conceitual dos dados do projeto.", 6.2)

add_para(anchor, "16.5 Relação entre modelagem, requisitos e rastreabilidade", "h2")
add_para(
    anchor,
    "A modelagem proposta fortalece a matriz de rastreabilidade porque associa requisitos funcionais, regras de "
    "negócio, atores e evidências visuais. Essa relação auxilia a verificar se funcionalidades críticas possuem "
    "representação no fluxo da solução e se há coerência entre requisitos, casos de uso, arquitetura conceitual e critérios de validação.",
)
add_table(
    anchor,
    "Tabela 12 – Complemento da rastreabilidade com casos de uso e diagramas",
    [
        ["Requisito", "Caso de uso relacionado", "Evidência de modelagem"],
        ["RF02 – Login e autenticação", "UC01", "Diagrama de atividades e sequência de login; classes User, Role, Permission e AuditLog."],
        ["RF03 – Cadastro de pacientes", "UC03", "Sequência de gerenciamento de paciente e entidade Patient."],
        ["RF05 – Registro de feridas", "UC04", "Classe Lesion, fluxo de avaliação e entidade Ferida."],
        ["RF06 – Upload de imagens", "UC05", "Classe ClinicalImage e vínculo com Ferida."],
        ["RF07 – Marcação de ROI", "UC06", "Classe ROI, atividade de ROI e análise assistiva."],
        ["RF10 – Relatório de acompanhamento", "UC11", "Classe Report e sequência de ROI, IA e relatório."],
        ["RF14 – Controle de permissões", "UC02", "Relação entre User, Role e Permission."],
        ["RF16 – Testes com bases públicas", "UC14", "ResearchExperiment e separação entre uso experimental e assistencial."],
        ["RF17 – Registro de métricas de IA", "UC14, UC15", "InferenceResult e ResearchExperiment."],
        ["RF19 – Log de auditoria", "UC01, UC02, UC19", "Classe AuditLog e fluxo de autenticação."],
        ["RF20 – Anonimização para pesquisa", "UC18", "Caso de uso de anonimização e preparação ética para pesquisa."],
    ],
)


replacements = {
    "16 TESTES COM BANCOS DE IMAGENS PÚBLICOS": "17 TESTES COM BANCOS DE IMAGENS PÚBLICOS",
    "17 TESTES DE REDIMENSIONAMENTO DE IMAGENS PARA IA": "18 TESTES DE REDIMENSIONAMENTO DE IMAGENS PARA IA",
    "18 USO DE ROI": "19 USO DE ROI",
    "19 INTELIGÊNCIA ARTIFICIAL E VISÃO COMPUTACIONAL": "20 INTELIGÊNCIA ARTIFICIAL E VISÃO COMPUTACIONAL",
    "20 ASPECTOS ÉTICOS": "21 ASPECTOS ÉTICOS",
    "21 LGPD E PROTEÇÃO DE DADOS": "22 LGPD E PROTEÇÃO DE DADOS",
    "22 PREPARAÇÃO PARA SUBMISSÃO AO CEP": "23 PREPARAÇÃO PARA SUBMISSÃO AO CEP",
    "23 RISCOS": "24 RISCOS",
    "24 CRITÉRIOS DE VALIDAÇÃO": "25 CRITÉRIOS DE VALIDAÇÃO",
    "25 CRONOGRAMA": "26 CRONOGRAMA",
    "26 CONSIDERAÇÕES FINAIS": "27 CONSIDERAÇÕES FINAIS",
    "Figura 2 – Representação metodológica do redimensionamento e registro de evidências": "Figura 11 – Representação metodológica do redimensionamento e registro de evidências",
    "Figura 3 – Fluxo seguro para ROI, análise assistiva e relatório": "Figura 12 – Fluxo seguro para ROI, análise assistiva e relatório",
    "Tabela 9 – Planejamento de testes com imagens públicas": "Tabela 13 – Planejamento de testes com imagens públicas",
    "Tabela 10 – Síntese interpretativa do redimensionamento": "Tabela 14 – Síntese interpretativa do redimensionamento",
    "Tabela 11 – Execução inicial dos filtros OpenCV em imagem sintética": "Tabela 15 – Execução inicial dos filtros OpenCV em imagem sintética",
    "Tabela 12 – Riscos do projeto": "Tabela 16 – Riscos do projeto",
    "Tabela 13 – Critérios de validação": "Tabela 17 – Critérios de validação",
    "Tabela 14 – Cronograma proposto": "Tabela 18 – Cronograma proposto",
    "Tabela 15 – Dicionário mínimo de dados": "Tabela 19 – Dicionário mínimo de dados",
}

for paragraph in doc.paragraphs:
    raw = paragraph.text.strip()
    if raw in replacements:
        is_heading = raw[0].isdigit() and "–" not in raw and not raw.startswith(("Tabela", "Figura"))
        set_text(paragraph, replacements[raw], 12, bold=is_heading, align=WD_ALIGN_PARAGRAPH.LEFT)


all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
if "GUEDES, Gilleanes" not in all_text:
    pressman = find_para(lambda text: text.startswith("PRESSMAN,"))
    paragraph = pressman.insert_paragraph_before(
        "GUEDES, Gilleanes T. A. UML 2: uma abordagem prática. 3. ed. São Paulo: Novatec, 2018."
    )
    format_body(paragraph)

if "WAZLAWICK, Raul" not in all_text:
    world = find_para(lambda text: text.startswith("WORLD COUNCIL") or text.startswith("WORLD HEALTH"))
    paragraph = world.insert_paragraph_before(
        "WAZLAWICK, Raul Sidnei. Engenharia de software: conceitos e práticas. Rio de Janeiro: Elsevier, 2013."
    )
    format_body(paragraph)


doc.save(OUTPUT_DOCX)
doc.save(OUTPUT_DOWNLOADS)
print(OUTPUT_DOCX)
print(OUTPUT_DOWNLOADS)
