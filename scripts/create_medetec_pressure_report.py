from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path

import cv2
from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOC_PATH = ROOT / "docs" / "research" / "relatorio_redisus_resolucao_com_imagens.docx"
REPORT_PATH = ROOT / "docs" / "research" / "relatorio_medetec_pressao_resolucao_nativa_vs_224.docx"
VALIDATION_SCRIPT = ROOT / "scripts" / "validate_medetec_pressure_resolution.py"
VALIDATION_OUTPUT = ROOT / "output" / "validation" / "medetec_pressure_resolution_validation.json"
CASE_IMAGE_DIR = ROOT / "tmp_images" / "medetec_pressure_cases"

if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.processing.clinical_wound_analyzer_core import ClinicalWoundAnalyzer


def copy_run_format(source_run, target_run) -> None:
    if source_run is None:
        return
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.style = source_run.style
    if source_run.font is not None and target_run.font is not None:
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.color.rgb = source_run.font.color.rgb
        target_run.font.highlight_color = source_run.font.highlight_color
        target_run.font.superscript = source_run.font.superscript
        target_run.font.subscript = source_run.font.subscript
        target_run.font.all_caps = source_run.font.all_caps
        target_run.font.small_caps = source_run.font.small_caps


def replace_paragraph_text(paragraph, text: str) -> None:
    source_run = paragraph.runs[0] if paragraph.runs else None
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("pPr"):
            continue
        p.remove(child)
    new_run = paragraph.add_run(text)
    copy_run_format(source_run, new_run)


def clear_cell(cell) -> None:
    cell.text = ""


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    source_run = paragraph.runs[0] if paragraph.runs else None
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    new_run = paragraph.add_run(text)
    copy_run_format(source_run, new_run)


def set_cell_image(cell, image_path: Path, width: float = 1.55) -> None:
    clear_cell(cell)
    run = cell.paragraphs[0].add_run()
    run.add_picture(str(image_path), width=Inches(width))


def ensure_validation_output() -> dict:
    if not VALIDATION_OUTPUT.exists():
        cmd = [
            sys.executable,
            str(VALIDATION_SCRIPT),
            "--output",
            str(VALIDATION_OUTPUT.relative_to(ROOT)),
        ]
        subprocess.run(cmd, cwd=ROOT, check=True)
    return json.loads(VALIDATION_OUTPUT.read_text(encoding="utf-8"))


def format_percent(rate: float) -> str:
    return f"{rate * 100:.2f}% ({rate:.4f})"


def format_stage(stage_code: str) -> str:
    if not stage_code:
        return "desconhecido"
    return stage_code.replace("stage_", "estÃ¡gio ")


def format_stage_counts(stage_counts: dict[str, int]) -> str:
    ordered = [f"{format_stage(code)}={count}" for code, count in sorted(stage_counts.items())]
    return ", ".join(ordered)


def format_tissue_counts(tissue_counts: dict[str, int]) -> str:
    ordered = [f"{name}={count}" for name, count in sorted(tissue_counts.items())]
    return ", ".join(ordered)


def select_cases(records: list[dict]) -> list[dict]:
    stable_candidates = [
        record
        for record in records
        if not record["pairwise"]["primary_tissue_changed"] and not record["pairwise"]["predicted_stage_changed"]
    ]
    stable = min(
        stable_candidates,
        key=lambda record: (
            float(record["pairwise"]["abs_health_score_delta"]),
            float(record["pairwise"]["abs_wound_fraction_delta"]),
        ),
    )

    tissue_shift_candidates = [
        record
        for record in records
        if record["pairwise"]["primary_tissue_changed"] and not record["pairwise"]["predicted_stage_changed"]
    ]
    tissue_shift = max(
        tissue_shift_candidates,
        key=lambda record: (
            float(record["pairwise"]["abs_health_score_delta"]),
            float(record["pairwise"]["abs_wound_fraction_delta"]),
        ),
    )

    stage_shift_candidates = [
        record
        for record in records
        if record["pairwise"]["predicted_stage_changed"] and not record["pairwise"]["primary_tissue_changed"]
    ]
    stage_shift = max(
        stage_shift_candidates,
        key=lambda record: (
            float(record["pairwise"]["abs_health_score_delta"]),
            float(record["pairwise"]["abs_wound_fraction_delta"]),
        ),
    )
    return [stable, tissue_shift, stage_shift]


def build_case_assets(case_records: list[dict]) -> list[dict]:
    CASE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    analyzer = ClinicalWoundAnalyzer()
    assets: list[dict] = []

    case_labels = [
        "Caso de Estudo 1: saÃ­da estÃ¡vel entre nativo e 224 x 224",
        "Caso de Estudo 2: mudanÃ§a de tecido primÃ¡rio apÃ³s resize para 224 x 224",
        "Caso de Estudo 3: mudanÃ§a de estÃ¡gio previsto apÃ³s resize para 224 x 224",
    ]

    for index, record in enumerate(case_records, start=1):
        image_path = Path(str(record["path"]))
        image = cv2.imread(str(image_path))
        if image is None:
            raise FileNotFoundError(f"NÃ£o foi possÃ­vel carregar {image_path}")

        resized = cv2.resize(image, (224, 224), interpolation=cv2.INTER_AREA)
        native_report = analyzer.analyze(image)
        resized_report = analyzer.analyze(resized)

        native_input_path = CASE_IMAGE_DIR / f"case{index}_native_input.png"
        native_detection_path = CASE_IMAGE_DIR / f"case{index}_native_detection.png"
        native_tissue_path = CASE_IMAGE_DIR / f"case{index}_native_tissue.png"
        resized_input_path = CASE_IMAGE_DIR / f"case{index}_224_input.png"
        resized_detection_path = CASE_IMAGE_DIR / f"case{index}_224_detection.png"
        resized_tissue_path = CASE_IMAGE_DIR / f"case{index}_224_tissue.png"

        cv2.imwrite(str(native_input_path), image)
        cv2.imwrite(str(native_detection_path), native_report.detection_overlay)
        cv2.imwrite(str(native_tissue_path), native_report.tissue_overlay)
        cv2.imwrite(str(resized_input_path), resized)
        cv2.imwrite(str(resized_detection_path), resized_report.detection_overlay)
        cv2.imwrite(str(resized_tissue_path), resized_report.tissue_overlay)

        pairwise = record["pairwise"]
        summary = (
            f"Arquivo: {image_path.name}. Tamanho nativo: {record['native_width']} x {record['native_height']}. "
            f"Pasta: {record['folder']}. "
            f"Nativo: tecido {record['native']['analyzer_primary_tissue']}, estÃ¡gio previsto {format_stage(record['native']['predicted_stage'])}, "
            f"health score {record['native']['analyzer_health_score']:.2f}, fraÃ§Ã£o da ferida {record['native']['analyzer_wound_fraction'] * 100:.2f}%. "
            f"224 x 224: tecido {record['resized_224']['analyzer_primary_tissue']}, estÃ¡gio previsto {format_stage(record['resized_224']['predicted_stage'])}, "
            f"health score {record['resized_224']['analyzer_health_score']:.2f}, fraÃ§Ã£o da ferida {record['resized_224']['analyzer_wound_fraction'] * 100:.2f}%. "
            f"Delta absoluto de health score: {pairwise['abs_health_score_delta']:.2f}. "
            f"Delta absoluto de fraÃ§Ã£o da Ã¡rea: {pairwise['abs_wound_fraction_delta'] * 100:.2f} pontos percentuais."
        )

        assets.append(
            {
                "record": record,
                "title": case_labels[index - 1],
                "summary": summary,
                "native_input_path": native_input_path,
                "native_detection_path": native_detection_path,
                "native_tissue_path": native_tissue_path,
                "resized_input_path": resized_input_path,
                "resized_detection_path": resized_detection_path,
                "resized_tissue_path": resized_tissue_path,
            }
        )

    return assets


def update_tables(doc: Document, case_assets: list[dict], payload: dict) -> None:
    if len(doc.tables) < 4:
        raise RuntimeError("O documento nÃ£o possui as tabelas esperadas.")

    native = payload["summary"]["native"]
    resized = payload["summary"]["resized_224"]
    pairwise = payload["summary"]["pairwise"]
    dataset = payload["dataset"]

    pipeline_table = doc.tables[0]
    table_rows = [
        [
            "MÃ©trica",
            "Medetec nativo",
            "Medetec 224 x 224",
            "Leitura tÃ©cnica",
            "ImplicaÃ§Ã£o",
        ],
        [
            "Base e resoluÃ§Ãµes",
            f"{dataset['total_images']} imagens Ãºteis; {dataset['unique_sizes']} tamanhos; {dataset['min_size']['width']} x {dataset['min_size']['height']} a {dataset['max_size']['width']} x {dataset['max_size']['height']}",
            "mesmas 175 imagens apÃ³s resize direto para 224 x 224 com cv2.INTER_AREA",
            "pressure_ulcers_1 e pressure_ulcers_2 foram usados; 2 introslides foram excluÃ­das; as pastas set_1/set_2 foram descartadas por duplicidade hash",
            "o acervo realmente permite estudar alta resoluÃ§Ã£o nativa versus reduÃ§Ã£o quadrada controlada",
        ],
        [
            "Processamento",
            f"sucesso {format_percent(native['analyzer_success_rate'])}; tempo mÃ©dio {native['analyzer_avg_processing_time_ms']:.2f} ms",
            f"sucesso {format_percent(resized['analyzer_success_rate'])}; tempo mÃ©dio {resized['analyzer_avg_processing_time_ms']:.2f} ms",
            "o ramo clÃ­nico nativo nÃ£o faz reduÃ§Ã£o interna porque o maior lado jÃ¡ Ã© menor que 1024 px",
            "224 x 224 ficou mais rÃ¡pido, porÃ©m velocidade nÃ£o significou equivalÃªncia clÃ­nica",
        ],
        [
            "Ferida vÃ¡lida",
            format_percent(native["analyzer_valid_wound_rate"]),
            format_percent(resized["analyzer_valid_wound_rate"]),
            f"acordo entre variantes: {format_percent(pairwise['valid_wound_agreement_rate'])}",
            "o resize nÃ£o derrubou a detecÃ§Ã£o de ferida, mas alterou leituras subsequentes",
        ],
        [
            "Tecido primÃ¡rio",
            format_tissue_counts(native["primary_tissue_counts"]),
            format_tissue_counts(resized["primary_tissue_counts"]),
            f"acordo {format_percent(pairwise['primary_tissue_agreement_rate'])}; {pairwise['primary_tissue_changed_count']} de {pairwise['total_pairs']} casos mudaram o tecido principal",
            "o resize para 224 x 224 aumentou a frequÃªncia de granulaÃ§Ã£o e reduziu esfacelo",
        ],
        [
            "EstÃ¡gio previsto",
            format_stage_counts(native["predicted_stage_counts"]),
            format_stage_counts(resized["predicted_stage_counts"]),
            f"acordo {format_percent(pairwise['predicted_stage_agreement_rate'])}; {pairwise['predicted_stage_changed_count']} de {pairwise['total_pairs']} casos mudaram o estÃ¡gio previsto",
            "na Medetec isso Ã© leitura comparativa, nÃ£o acurÃ¡cia, porque o acervo nÃ£o traz rÃ³tulo de estÃ¡gio",
        ],
        [
            "ROI e health score",
            f"fraÃ§Ã£o mÃ©dia da Ã¡rea {native['avg_wound_fraction'] * 100:.2f}%; health mÃ©dio {native['avg_health_score']:.2f}",
            f"fraÃ§Ã£o mÃ©dia da Ã¡rea {resized['avg_wound_fraction'] * 100:.2f}%; health mÃ©dio {resized['avg_health_score']:.2f}",
            f"delta mÃ©dio absoluto: Ã¡rea {pairwise['avg_abs_wound_fraction_delta'] * 100:.2f} pp; health {pairwise['avg_abs_health_score_delta']:.2f}",
            "o resize quadrado alterou tamanho relativo da ROI e a leitura de severidade/saÃºde",
        ],
        [
            "SÃ­ntese",
            "mais detalhe espacial preservado, maior predominÃ¢ncia de esfacelo e health score mais baixo",
            "mesma capacidade de processamento, porÃ©m com tendÃªncia a suavizar padrÃµes e elevar granulaÃ§Ã£o/health score",
            "os casos reais abaixo mostram estabilidade, mudanÃ§a de tecido e mudanÃ§a de estÃ¡gio previsto",
            "para relatÃ³rio clÃ­nico, a imagem nativa Ã© mais fiel que o achatamento para 224 x 224",
        ],
    ]

    for row_idx, row_values in enumerate(table_rows):
        row = pipeline_table.rows[row_idx]
        for col_idx, value in enumerate(row_values):
            set_cell_text(row.cells[col_idx], value)

    row_labels = [
        "Entrada nativa",
        "ROI / detecÃ§Ã£o nativa",
        "Tecidos nativos",
        "Entrada 224 x 224",
        "ROI / detecÃ§Ã£o 224 x 224",
        "Tecidos 224 x 224",
    ]

    for table_index, case in enumerate(case_assets, start=1):
        table = doc.tables[table_index]
        while len(table.rows) < 7:
            table.add_row()

        header = ["Etapa", "SaÃ­da visual", "DescriÃ§Ã£o tÃ©cnica"]
        for col_idx, value in enumerate(header):
            set_cell_text(table.rows[0].cells[col_idx], value)

        record = case["record"]
        native = record["native"]
        resized = record["resized_224"]
        descriptions = [
            (
                f"Imagem original {record['native_width']} x {record['native_height']}. "
                "No pipeline clÃ­nico, esta variante entra sem reduÃ§Ã£o proporcional automÃ¡tica porque o maior lado Ã© menor que 1024 px."
            ),
            (
                f"Ferida vÃ¡lida: sim. FraÃ§Ã£o da Ã¡rea: {native['analyzer_wound_fraction'] * 100:.2f}%. "
                f"Health score: {native['analyzer_health_score']:.2f}. EstÃ¡gio previsto: {format_stage(native['predicted_stage'])}."
            ),
            (
                f"Tecido primÃ¡rio: {native['analyzer_primary_tissue']}. "
                f"Tempo de processamento: {native['analyzer_processing_time_ms']:.2f} ms."
            ),
            "A mesma imagem foi comprimida diretamente para 224 x 224 com cv2.INTER_AREA antes da anÃ¡lise.",
            (
                f"Ferida vÃ¡lida: sim. FraÃ§Ã£o da Ã¡rea: {resized['analyzer_wound_fraction'] * 100:.2f}%. "
                f"Health score: {resized['analyzer_health_score']:.2f}. EstÃ¡gio previsto: {format_stage(resized['predicted_stage'])}."
            ),
            (
                f"Tecido primÃ¡rio: {resized['analyzer_primary_tissue']}. "
                f"Tempo de processamento: {resized['analyzer_processing_time_ms']:.2f} ms."
            ),
        ]
        image_paths = [
            case["native_input_path"],
            case["native_detection_path"],
            case["native_tissue_path"],
            case["resized_input_path"],
            case["resized_detection_path"],
            case["resized_tissue_path"],
        ]

        for row_offset in range(6):
            row = table.rows[row_offset + 1]
            set_cell_text(row.cells[0], row_labels[row_offset])
            set_cell_image(row.cells[1], image_paths[row_offset], width=1.6)
            set_cell_text(row.cells[2], descriptions[row_offset])


def update_paragraphs(doc: Document, case_assets: list[dict], payload: dict) -> None:
    paragraphs = doc.paragraphs
    dataset = payload["dataset"]
    native = payload["summary"]["native"]
    resized = payload["summary"]["resized_224"]
    pairwise = payload["summary"]["pairwise"]

    replace_paragraph_text(
        paragraphs[12],
        "RelatÃ³rio TÃ©cnico: Medetec de LesÃ£o por PressÃ£o em ResoluÃ§Ã£o Nativa versus 224 x 224",
    )
    replace_paragraph_text(paragraphs[29], "16 de abril de 2026")
    replace_paragraph_text(
        paragraphs[39],
        "RelatÃ³rio TÃ©cnico: Medetec de LesÃ£o por PressÃ£o em ResoluÃ§Ã£o Nativa versus 224 x 224",
    )

    replace_paragraph_text(
        paragraphs[92],
        (
            "Este documento aplica o mesmo racional do relatÃ³rio anterior ao subconjunto Medetec de lesÃ£o por pressÃ£o, "
            "mas agora com imagens nativas maiores que 224 x 224 e com comparaÃ§Ã£o real entre a captura original e a mesma "
            "imagem reduzida diretamente para 224 x 224. Foram analisadas 175 imagens Ãºteis das pastas pressure_ulcers_1 "
            "e pressure_ulcers_2, apÃ³s exclusÃ£o de 2 introslides e descarte das pastas espelhadas set_1_of_2 e set_2_of_2 "
            "por duplicidade exata de hash. O acervo nativo apresentou 79 resoluÃ§Ãµes, variando de "
            f"{dataset['min_size']['width']} x {dataset['min_size']['height']} a {dataset['max_size']['width']} x {dataset['max_size']['height']}. "
            f"O HEAL+ processou 175 de 175 imagens com sucesso tanto no nativo quanto no 224 x 224 e aceitou 100% dos casos como ferida vÃ¡lida em ambas as variantes. "
            f"Apesar disso, o resize quadrado para 224 x 224 alterou o tecido primÃ¡rio em {pairwise['primary_tissue_changed_count']} de 175 casos "
            f"({format_percent(pairwise['primary_tissue_changed_count'] / 175)}) e mudou o estÃ¡gio previsto em {pairwise['predicted_stage_changed_count']} de 175 casos "
            f"({format_percent(pairwise['predicted_stage_changed_count'] / 175)}). O acordo entre variantes foi de "
            f"{format_percent(pairwise['primary_tissue_agreement_rate'])} para tecido primÃ¡rio e {format_percent(pairwise['predicted_stage_agreement_rate'])} para estÃ¡gio previsto. "
            f"O health score mÃ©dio subiu de {native['avg_health_score']:.2f} no nativo para {resized['avg_health_score']:.2f} no 224 x 224, "
            f"com delta absoluto mÃ©dio de {pairwise['avg_abs_health_score_delta']:.2f} pontos."
        ),
    )
    replace_paragraph_text(
        paragraphs[93],
        "Palavras-chave: Medetec, lesÃ£o por pressÃ£o, resoluÃ§Ã£o nativa, 224 x 224, ROI, HEAL+.",
    )

    replace_paragraph_text(paragraphs[97], "1 OBJETIVO")
    replace_paragraph_text(
        paragraphs[98],
        "Documentar, no subconjunto Medetec de lesÃ£o por pressÃ£o, como o pipeline real HEAL+ lÃª a imagem nativa maior, como a ROI Ã© extraÃ­da, como contraste e segmentaÃ§Ã£o se comportam e o que muda quando a mesma imagem Ã© redimensionada diretamente para 224 x 224.",
    )

    replace_paragraph_text(paragraphs[99], "2 RECORTE EXPERIMENTAL")
    replace_paragraph_text(
        paragraphs[100],
        "O recorte foi executado exclusivamente sobre as imagens de lesÃ£o por pressÃ£o da Medetec, com o seguinte desenho de base:",
    )
    replace_paragraph_text(paragraphs[101], "pressure_ulcers_1: 101 imagens Ãºteis apÃ³s excluir 1 introslide;")
    replace_paragraph_text(paragraphs[102], "pressure_ulcers_2: 74 imagens Ãºteis apÃ³s excluir 1 introslide;")
    replace_paragraph_text(paragraphs[103], "pressure_ulcers_set_1_of_2 e pressure_ulcers_set_2_of_2 existem no acervo local, porÃ©m foram excluÃ­das por serem cÃ³pias hash-idÃªnticas das pastas 1 e 2;")
    replace_paragraph_text(paragraphs[104], f"total final: {dataset['total_images']} imagens Ãºnicas, com {dataset['unique_sizes']} resoluÃ§Ãµes nativas entre {dataset['min_size']['width']} x {dataset['min_size']['height']} e {dataset['max_size']['width']} x {dataset['max_size']['height']}.")
    replace_paragraph_text(
        paragraphs[106],
        "Diferentemente do PIID, este acervo realmente oferece imagens maiores que 224 x 224, permitindo comparar a anÃ¡lise na resoluÃ§Ã£o nativa com a mesma imagem comprimida para 224 x 224 sem recorrer a um experimento puramente sintÃ©tico de multi-resoluÃ§Ã£o.",
    )

    replace_paragraph_text(paragraphs[107], "3 PIPELINE AVALIADO")
    replace_paragraph_text(paragraphs[108], "A avaliaÃ§Ã£o real considerou o fluxo efetivamente executado hoje pelo HEAL+ no motor clÃ­nico headless.")
    replace_paragraph_text(paragraphs[110], "ClinicalWoundAnalyzer e heal_analyzer.py como nÃºcleo do pipeline clÃ­nico real;")
    replace_paragraph_text(paragraphs[112], "ImageEnhancer, WoundDetectorCV e ROISegmenter para correÃ§Ã£o de qualidade, detecÃ§Ã£o inicial, mÃ¡scara por contorno e zonas espaciais;")
    replace_paragraph_text(paragraphs[115], "_segment_clinical_v3 e PressureInjuryStageClassifier para leitura tecidual, health score e estÃ¡gio previsto.")
    replace_paragraph_text(
        paragraphs[117],
        "A comparaÃ§Ã£o nÃ£o mede acurÃ¡cia clÃ­nica contra rÃ³tulo Medetec, porque esse acervo nÃ£o fornece estÃ¡gio verdadeiro; ela mede estabilidade do pipeline entre a variante nativa e a variante 224 x 224 da mesma imagem.",
    )

    replace_paragraph_text(paragraphs[119], "4 METODOLOGIA")
    replace_paragraph_text(
        paragraphs[120],
        "Cada imagem foi processada em dois cenÃ¡rios, mantendo o restante do pipeline idÃªntico.",
    )
    replace_paragraph_text(paragraphs[121], "ler a imagem nativa da Medetec com cv2.imread, preservando a geometria original em torno de 560 x 4xx ou 560 x 560;")
    replace_paragraph_text(paragraphs[122], "gerar uma segunda variante da mesma imagem via cv2.resize(..., (224, 224), interpolation=cv2.INTER_AREA), sem padding e sem preservaÃ§Ã£o explÃ­cita da razÃ£o de aspecto;")
    replace_paragraph_text(paragraphs[123], "rodar ClinicalWoundAnalyzer.analyze nas duas variantes e registrar validade, tecido primÃ¡rio, fraÃ§Ã£o da Ã¡rea, health score e tempo de processamento;")
    replace_paragraph_text(paragraphs[124], "rodar PressureInjuryStageClassifier nas duas variantes para observar estabilidade do estÃ¡gio previsto sob resize direto;")
    replace_paragraph_text(paragraphs[125], "selecionar casos reais para ilustrar trÃªs cenÃ¡rios: saÃ­da estÃ¡vel, mudanÃ§a de tecido e mudanÃ§a de estÃ¡gio previsto.")
    replace_paragraph_text(paragraphs[127], "Foram consolidados os seguintes indicadores comparativos:")
    replace_paragraph_text(paragraphs[128], "taxa de leitura e processamento em nativo e em 224 x 224;")
    replace_paragraph_text(paragraphs[129], "taxa de ferida vÃ¡lida em ambas as variantes;")
    replace_paragraph_text(paragraphs[130], "acordo de tecido primÃ¡rio e acordo de estÃ¡gio previsto entre a imagem nativa e a versÃ£o 224 x 224;")
    replace_paragraph_text(paragraphs[131], "delta absoluto mÃ©dio de fraÃ§Ã£o da Ã¡rea da ferida e de health score;")
    replace_paragraph_text(paragraphs[132], "redistribuiÃ§Ã£o das contagens de tecido primÃ¡rio e de estÃ¡gio previsto apÃ³s o resize;")
    replace_paragraph_text(paragraphs[133], "estudos de caso reais com visualizaÃ§Ã£o da entrada, ROI e overlay tecidual nos dois cenÃ¡rios.")

    replace_paragraph_text(paragraphs[134], "5 TRATAMENTO DE RESOLUÃ‡ÃƒO E DESENHO COMPARATIVO")
    replace_paragraph_text(
        paragraphs[135],
        "No subconjunto Medetec avaliado, a diferenÃ§a entre as variantes Ã© controlada e tecnicamente rastreÃ¡vel.",
    )
    replace_paragraph_text(
        paragraphs[136],
        "Na variante nativa, ClinicalWoundAnalyzer._prepare_input nÃ£o reduz a imagem porque o maior lado jÃ¡ Ã© inferior a 1024 px; assim, a geometria original da Medetec Ã© preservada.",
    )
    replace_paragraph_text(
        paragraphs[137],
        "Na variante comparativa, a imagem Ã© comprimida externamente para 224 x 224 por resize quadrado direto, o que altera a razÃ£o de aspecto quando a origem nÃ£o Ã© quadrada.",
    )
    replace_paragraph_text(
        paragraphs[138],
        "Esse cenÃ¡rio emula de forma realista o que acontece quando um fluxo prÃ©vio forÃ§a compatibilidade com entradas quadradas antes do pipeline clÃ­nico.",
    )
    replace_paragraph_text(
        paragraphs[139],
        "A comparaÃ§Ã£o, portanto, testa nÃ£o apenas perda de detalhe espacial, mas tambÃ©m a distorÃ§Ã£o geomÃ©trica introduzida pelo achatamento para 224 x 224.",
    )

    replace_paragraph_text(paragraphs[140], "6 CONTEXTO DA BASE MEDETEC")
    replace_paragraph_text(
        paragraphs[141],
        "O subconjunto Medetec de pressÃ£o nÃ£o possui mÃ¡scaras manuais, Ã¡reas anotadas nem estÃ¡gio clÃ­nico verdadeiro no acervo local; ele funciona aqui como base externa para avaliar a sensibilidade do pipeline Ã  resoluÃ§Ã£o e ao resize.",
    )
    replace_paragraph_text(
        paragraphs[142],
        "A base tambÃ©m contÃ©m duplicatas de organizaÃ§Ã£o: pressure_ulcers_set_1_of_2 replica pressure_ulcers_1 e pressure_ulcers_set_2_of_2 replica pressure_ulcers_2. Esses espelhos foram removidos para evitar dupla contagem.",
    )
    replace_paragraph_text(
        paragraphs[143],
        "Como a origem nativa jÃ¡ estÃ¡ majoritariamente em 560 x 4xx, esta anÃ¡lise preenche justamente a lacuna deixada pelo PIID, onde o acervo local jÃ¡ nasce em 224 x 224.",
    )

    replace_paragraph_text(paragraphs[145], "7 RESULTADOS CONSOLIDADOS")
    replace_paragraph_text(paragraphs[146], "Tabela 1 - Resumo comparativo da Medetec nativa versus a mesma imagem em 224 x 224")
    replace_paragraph_text(
        paragraphs[148],
        "Fonte: elaboraÃ§Ã£o prÃ³pria com base em dataset/medetec/metadata.json, validate_medetec_pressure_resolution.py e output/validation/medetec_pressure_resolution_validation.json.",
    )
    replace_paragraph_text(paragraphs[149], "Na comparaÃ§Ã£o real entre as duas variantes, o comportamento consolidado foi o seguinte:")
    replace_paragraph_text(
        paragraphs[150],
        f"o HEAL+ processou {native['total_images']} de {native['total_images']} imagens com sucesso tanto no nativo quanto no 224 x 224, e aceitou 100% dos casos como ferida vÃ¡lida em ambas as leituras;",
    )
    replace_paragraph_text(
        paragraphs[151],
        f"o resize para 224 x 224 mudou o tecido primÃ¡rio em {pairwise['primary_tissue_changed_count']} de {pairwise['total_pairs']} casos e elevou o health score mÃ©dio de {native['avg_health_score']:.2f} para {resized['avg_health_score']:.2f};",
    )
    replace_paragraph_text(
        paragraphs[152],
        f"o estÃ¡gio previsto mudou em {pairwise['predicted_stage_changed_count']} de {pairwise['total_pairs']} casos, enquanto a concordÃ¢ncia geral foi de {format_percent(pairwise['predicted_stage_agreement_rate'])}; como nÃ£o hÃ¡ rÃ³tulo Medetec, isso deve ser interpretado como estabilidade relativa, nÃ£o como acurÃ¡cia.",
    )

    replace_paragraph_text(paragraphs[153], "7.1 Estudos de caso reais: nativo versus 224 x 224")
    replace_paragraph_text(
        paragraphs[154],
        "Os estudos de caso abaixo usam exatamente a mesma imagem Medetec em dois cenÃ¡rios: a entrada nativa maior e a mesma entrada comprimida para 224 x 224. Cada tabela mostra entrada, detecÃ§Ã£o/ROI e overlay tecidual nas duas variantes.",
    )

    case_heading_indices = [156, 166, 171]
    case_summary_indices = [158, 168, 173]
    for idx, case in enumerate(case_assets):
        replace_paragraph_text(paragraphs[case_heading_indices[idx]], case["title"])
        replace_paragraph_text(paragraphs[case_summary_indices[idx]], case["summary"])

    replace_paragraph_text(paragraphs[174], "8 IMPACTO DO REDIMENSIONAMENTO")
    replace_paragraph_text(paragraphs[175], "Na leitura comparativa entre a Medetec nativa e a mesma imagem em 224 x 224, os principais achados foram:")
    replace_paragraph_text(
        paragraphs[177],
        "a validaÃ§Ã£o de ferida permaneceu estÃ¡vel em 100% dos pares, indicando que o resize nÃ£o fez o analisador deixar de reconhecer lesÃµes;",
    )
    replace_paragraph_text(
        paragraphs[178],
        f"o tecido primÃ¡rio mudou em {pairwise['primary_tissue_changed_count']} de {pairwise['total_pairs']} pares, com aumento marcante de Tecido de GranulaÃ§Ã£o de {native['primary_tissue_counts'].get('Tecido de GranulaÃ§Ã£o', 0)} para {resized['primary_tissue_counts'].get('Tecido de GranulaÃ§Ã£o', 0)} casos;",
    )
    replace_paragraph_text(
        paragraphs[179],
        f"o estÃ¡gio previsto mudou em {pairwise['predicted_stage_changed_count']} de {pairwise['total_pairs']} pares, e a distribuiÃ§Ã£o deslocou mais casos para estÃ¡gio 4 apÃ³s o resize ({native['predicted_stage_counts'].get('stage_4', 0)} para {resized['predicted_stage_counts'].get('stage_4', 0)});",
    )
    replace_paragraph_text(
        paragraphs[180],
        f"o delta absoluto mÃ©dio foi de {pairwise['avg_abs_wound_fraction_delta'] * 100:.2f} pontos percentuais na fraÃ§Ã£o da ROI e de {pairwise['avg_abs_health_score_delta']:.2f} pontos no health score, sugerindo alteraÃ§Ã£o clÃ­nica relevante mesmo sem perda da taxa de processamento.",
    )

    replace_paragraph_text(paragraphs[182], "9 CONCLUSÃƒO E RECOMENDAÃ‡ÃƒO")
    replace_paragraph_text(
        paragraphs[183],
        "Com base na Medetec de lesÃ£o por pressÃ£o, Ã© correto afirmar que o HEAL+ consegue processar imagens nativas maiores e que o comportamento clÃ­nico nÃ£o Ã© idÃªntico ao da mesma imagem comprimida para 224 x 224. Nesta base externa, o resize direto preservou a taxa de ferida vÃ¡lida, mas alterou a interpretaÃ§Ã£o de tecido, Ã¡rea relativa, health score e, em parte dos casos, o estÃ¡gio previsto.",
    )
    replace_paragraph_text(paragraphs[185], f"o pipeline clÃ­nico atual leu e processou {native['total_images']}/{native['total_images']} imagens nas duas variantes;")
    replace_paragraph_text(paragraphs[187], "100% dos pares permaneceram vÃ¡lidos como ferida apÃ³s o resize, mas isso nÃ£o implicou equivalÃªncia clÃ­nica;")
    replace_paragraph_text(paragraphs[189], f"o health score mÃ©dio aumentou {resized['avg_health_score'] - native['avg_health_score']:.2f} pontos ao passar do nativo para 224 x 224;")
    replace_paragraph_text(paragraphs[191], f"o tecido primÃ¡rio mudou em {pairwise['primary_tissue_changed_count']} casos e o estÃ¡gio previsto mudou em {pairwise['predicted_stage_changed_count']} casos.")
    replace_paragraph_text(
        paragraphs[193],
        "Para lesÃ£o por pressÃ£o em imagens originalmente maiores, a recomendaÃ§Ã£o Ã© manter a resoluÃ§Ã£o nativa atÃ© o ponto em que o motor clÃ­nico possa trabalhar sem achatamento quadrado. Se algum ramo profundo exigir 224 x 224, o ideal Ã© usar letterbox ou padding, e nÃ£o distorÃ§Ã£o geomÃ©trica direta.",
    )
    replace_paragraph_text(paragraphs[194], "manter a imagem nativa da Medetec como referÃªncia clÃ­nica principal;")
    replace_paragraph_text(paragraphs[195], "evitar prÃ©-redimensionamento quadrado direto quando o objetivo for laudo tecidual e health score;")
    replace_paragraph_text(paragraphs[196], "priorizar padding ou letterbox nos ramos que ainda exigirem entrada fixa 224 x 224.")
    replace_paragraph_text(paragraphs[198], "Assim, a interpretaÃ§Ã£o final deste relatÃ³rio Ã© a seguinte:")
    replace_paragraph_text(paragraphs[200], "evidÃªncia real atual: a Medetec nativa maior preserva processamento estÃ¡vel e produz leituras clinicamente diferentes do 224 x 224;")
    replace_paragraph_text(paragraphs[201], "evidÃªncia ainda ausente: ground truth de estÃ¡gio, mÃ¡scaras manuais e medidas clÃ­nicas externas para validar qual variante estÃ¡ mais prÃ³xima do julgamento especialista;")
    replace_paragraph_text(paragraphs[202], "principal risco tÃ©cnico a revisar: resize quadrado direto antes da anÃ¡lise tecidual e da classificaÃ§Ã£o de estÃ¡gio.")

    replace_paragraph_text(paragraphs[203], "10 LIMITAÃ‡Ã•ES")
    replace_paragraph_text(paragraphs[204], "Este resultado deve ser interpretado com as seguintes ressalvas:")
    replace_paragraph_text(paragraphs[205], "a Medetec local nÃ£o traz rÃ³tulo verdadeiro de estÃ¡gio nem mÃ¡scaras manuais, entÃ£o a comparaÃ§Ã£o de estÃ¡gio Ã© apenas relativa entre variantes;")
    replace_paragraph_text(paragraphs[206], "o acervo Ã© heterogÃªneo, com iluminaÃ§Ã£o, enquadramento e contexto clÃ­nico variados por se tratar de base externa raspada da web;")
    replace_paragraph_text(paragraphs[207], "o ensemble adicional baseado em DermaIntel, MedSAM e BiomedCLIP permaneceu indisponÃ­vel no ambiente atual por dependÃªncias ausentes;")
    replace_paragraph_text(paragraphs[208], "esta anÃ¡lise testa resize direto para 224 x 224; outras estratÃ©gias, como padding, ainda precisam ser medidas separadamente.")

    replace_paragraph_text(paragraphs[209], "11 REPRODUTIBILIDADE")
    replace_paragraph_text(paragraphs[210], "A validaÃ§Ã£o comparativa descrita neste documento pode ser reproduzida pelo script:")
    replace_paragraph_text(paragraphs[211], "scripts/validate_medetec_pressure_resolution.py")
    replace_paragraph_text(paragraphs[212], "O resultado consolidado usado nesta revisÃ£o foi salvo em:")
    replace_paragraph_text(paragraphs[213], "output/validation/medetec_pressure_resolution_validation.json")

    replace_paragraph_text(paragraphs[214], "12 ENCAMINHAMENTO PRÁTICO")
    replace_paragraph_text(paragraphs[215], "Para produto, coleta e treinamento, recomenda-se adotar imediatamente:")
    replace_paragraph_text(paragraphs[216], "manter a maior resoluÃ§Ã£o nativa disponÃ­vel atÃ© a etapa clÃ­nica do pipeline;")
    replace_paragraph_text(paragraphs[217], "usar 224 x 224 apenas como requisito tÃ©cnico de compatibilidade, nÃ£o como substituto neutro da imagem maior;")
    replace_paragraph_text(paragraphs[218], "revisar os ramos profundos para receber letterbox ou padding e medir novamente a estabilidade na Medetec;")
    replace_paragraph_text(paragraphs[219], "se o objetivo for publicaÃ§Ã£o ou decisÃ£o clÃ­nica comparativa, incluir revisÃ£o especialista dos casos em que nativo e 224 x 224 divergiram.")

    replace_paragraph_text(paragraphs[236], "REDISUS. Medetec: metadados do subconjunto de lesÃ£o por pressÃ£o. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[237], "dataset/medetec/metadata.json. Acesso em: 16 abr. 2026.")
    replace_paragraph_text(paragraphs[239], "REDISUS. ValidaÃ§Ã£o comparativa Medetec nativa versus 224 x 224. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[240], "output/validation/medetec_pressure_resolution_validation.json. Acesso em: 16 abr. 2026.")
    replace_paragraph_text(paragraphs[242], "REDISUS. Script de validaÃ§Ã£o comparativa da Medetec em lesÃ£o por pressÃ£o. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[243], "scripts/validate_medetec_pressure_resolution.py. Acesso em: 16 abr. 2026.")


def create_report() -> Path:
    if not TEMPLATE_DOC_PATH.exists():
        raise FileNotFoundError(f"Modelo de relatÃ³rio nÃ£o encontrado: {TEMPLATE_DOC_PATH}")

    payload = ensure_validation_output()
    case_records = select_cases(payload["records"])
    case_assets = build_case_assets(case_records)

    target_path = REPORT_PATH
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        shutil.copy2(TEMPLATE_DOC_PATH, target_path)
    except PermissionError:
        target_path = REPORT_PATH.with_stem(REPORT_PATH.stem + "_atualizado")
        shutil.copy2(TEMPLATE_DOC_PATH, target_path)

    doc = Document(target_path)
    update_paragraphs(doc, case_assets, payload)
    update_tables(doc, case_assets, payload)
    doc.save(target_path)
    return target_path


def main() -> int:
    output_path = create_report()
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


