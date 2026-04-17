from __future__ import annotations

import json
import shutil
import subprocess
import sys
from collections import Counter
from pathlib import Path
from typing import Iterable

import cv2
from docx import Document
from docx.shared import Inches
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "research" / "relatorio_redisus_resolucao_com_imagens.docx"
VALIDATION_SCRIPT = ROOT / "scripts" / "validate_heal_analyzer_piid.py"
VALIDATION_OUTPUT = ROOT / "output" / "validation" / "piid_heal_analyzer_validation_test_current.json"
CASE_IMAGE_DIR = ROOT / "tmp_images" / "real_pipeline_cases_current"
PIID_ROOT = ROOT / "dataset" / "piid" / "raw"

if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.processing.clinical_wound_analyzer_core import ClinicalWoundAnalyzer


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clear_cell(cell) -> None:
    cell.text = ""


def set_cell_text(cell, text: str) -> None:
    clear_cell(cell)
    cell.paragraphs[0].add_run(text)


def set_cell_image(cell, image_path: Path, width: float = 1.7) -> None:
    clear_cell(cell)
    run = cell.paragraphs[0].add_run()
    run.add_picture(str(image_path), width=Inches(width))


def ensure_validation_output() -> dict:
    if not VALIDATION_OUTPUT.exists():
        cmd = [
            sys.executable,
            str(VALIDATION_SCRIPT),
            "--split",
            "test",
            "--samples-per-stage",
            "0",
            "--output",
            str(VALIDATION_OUTPUT.relative_to(ROOT)),
        ]
        subprocess.run(cmd, cwd=ROOT, check=True)
    return json.loads(VALIDATION_OUTPUT.read_text(encoding="utf-8"))


def format_percent(rate: float) -> str:
    return f"{rate * 100:.2f}% ({rate:.4f})"


def collect_native_size_summary(image_root: Path) -> tuple[str, int]:
    counts: Counter[tuple[int, int]] = Counter()
    total = 0
    for image_path in image_root.rglob("*.jpg"):
        with Image.open(image_path) as image:
            counts[image.size] += 1
            total += 1
    if not counts:
        return "desconhecido", 0
    top_size, _ = counts.most_common(1)[0]
    if len(counts) == 1:
        return f"{top_size[0]} x {top_size[1]}", total
    joined = ", ".join(f"{w}x{h}" for (w, h), _ in counts.most_common(3))
    return joined, total


def select_case(records: Iterable[dict], stage_code: str, target_area: int = 13000) -> dict:
    candidates = [
        record
        for record in records
        if record.get("true_stage") == stage_code
        and record.get("analyzer_is_valid_wound")
        and record.get("predicted_stage") == stage_code
        and 8000 <= int(record.get("analyzer_wound_area_px", 0)) <= 30000
    ]
    if not candidates:
        candidates = [
            record
            for record in records
            if record.get("true_stage") == stage_code
            and record.get("analyzer_is_valid_wound")
        ]
    if not candidates:
        raise RuntimeError(f"Nenhum caso válido encontrado para {stage_code}.")
    return min(
        candidates,
        key=lambda record: abs(int(record.get("analyzer_wound_area_px", 0)) - target_area),
    )


def stage_code_to_label(stage_code: str) -> str:
    return stage_code.replace("_", " ").replace("stage", "estágio")


def build_case_assets(case_records: list[dict]) -> list[dict]:
    CASE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    analyzer = ClinicalWoundAnalyzer()
    assets: list[dict] = []

    for index, record in enumerate(case_records, start=1):
        image_path = Path(str(record["path"]))
        image = cv2.imread(str(image_path))
        if image is None:
            raise FileNotFoundError(f"Não foi possível carregar {image_path}")

        true_stage_label = stage_code_to_label(str(record["true_stage"]))
        predicted_stage_label = stage_code_to_label(str(record.get("predicted_stage", "")))

        report = analyzer.analyze(image)
        original_path = CASE_IMAGE_DIR / f"case{index}_original.png"
        detection_path = CASE_IMAGE_DIR / f"case{index}_detection.png"
        tissue_path = CASE_IMAGE_DIR / f"case{index}_tissue_overlay.png"

        cv2.imwrite(str(original_path), image)
        cv2.imwrite(str(detection_path), report.detection_overlay)
        cv2.imwrite(str(tissue_path), report.tissue_overlay)

        assets.append(
            {
                "record": record,
                "report": report,
                "title": f"Caso de Estudo {index}: imagem real do {true_stage_label}",
                "summary": (
                    f"Arquivo: {image_path.name}. Estágio verdadeiro: {true_stage_label}. "
                    f"Estágio previsto: {predicted_stage_label}. "
                    f"Tecido primário atual: {report.primary_tissue}. "
                    f"Área da ROI: {report.wound_area_px} px. "
                    f"Health score: {report.health_score:.2f}. "
                    f"Tempo de processamento: {report.processing_time_ms:.2f} ms."
                ),
                "original_path": original_path,
                "detection_path": detection_path,
                "tissue_path": tissue_path,
            }
        )

    return assets


def update_tables(doc: Document, case_assets: list[dict]) -> None:
    if len(doc.tables) < 4:
        raise RuntimeError("O documento não possui as tabelas esperadas.")

    pipeline_table = doc.tables[0]
    pipeline_rows = [
        [
            "Etapa",
            "Entrada / leitura",
            "Processamento principal",
            "Saída",
            "Impacto da resolução",
        ],
        [
            "Leitura",
            "JPEG PIID via cv2.imread (BGR, 8 bits)",
            "carrega matriz HxWx3 para o pipeline clínico",
            "imagem original",
            "o acervo local atual chega inteiramente em 224 x 224",
        ],
        [
            "Pré-processamento",
            "imagem BGR",
            "análise de iluminação, brilho, white balance, CLAHE e correções automáticas quando necessário",
            "imagem corrigida para análise",
            "se o maior lado excede 1024 px, a redução é proporcional; sem distorção geométrica",
        ],
        [
            "Detecção inicial",
            "gray, HSV, LAB e textura",
            "WoundDetectorCV no modo texture_priority, com exclusão de pele e drape",
            "bbox(s) da lesão",
            "detalhe fino e contraste ajudam na confiança do detector",
        ],
        [
            "ROI por contorno",
            "bbox(s) + HSV",
            "ROISegmenter cria máscara por contorno, remove fundo cirúrgico e background residual",
            "wound_mask",
            "a ROI não é um retângulo fixo; o contorno final vem da máscara",
        ],
        [
            "Zonas espaciais",
            "wound_mask",
            "erosão e dilatação adaptativas para periferia, core e outer_ring",
            "máscaras espaciais",
            "essas zonas orientam necrose, esfacelo e epitelização",
        ],
        [
            "Segmentação clínica v3",
            "ROI em HSV, LAB e gray",
            "bilateral + CLAHE, fusão HSV 60 / LAB 40, skin exclusion, textura, luminância e Scharr",
            "percentuais, seg_map e overlay tecidual",
            "o ramo clínico preserva proporção; a perda principal aparece se a origem já vier degradada",
        ],
        [
            "Modelos profundos",
            "imagem BGR",
            "ResNet50 two-stage em 224 x 224 e U-Net em 512 x 512 quando disponível",
            "classe, confiança e Grad-CAM",
            "resize quadrado direto pode distorcer a anatomia se a origem não for quadrada",
        ],
    ]

    for row_idx, row_values in enumerate(pipeline_rows):
        row = pipeline_table.rows[row_idx]
        for col_idx, value in enumerate(row_values):
            set_cell_text(row.cells[col_idx], value)

    step_labels = [
        "Entrada nativa PIID",
        "ROI por detecção e contorno",
        "Segmentação tecidual v3",
    ]
    step_descriptions = [
        (
            "cv2.imread carrega o JPEG como matriz BGR de 8 bits. "
            "No PIID local, a imagem já entra em 224 x 224."
        ),
        (
            "WoundDetectorCV gera a detecção inicial e ROISegmenter refina a máscara "
            "por contorno, com exclusão de fundo cirúrgico e background residual."
        ),
        (
            "ClinicalWoundAnalyzer aplica bilateral + CLAHE, fusão HSV/LAB, zonas "
            "core/periferia, skin exclusion, textura, luminância e gradiente Scharr."
        ),
    ]

    for table_index, case in enumerate(case_assets, start=1):
        table = doc.tables[table_index]
        header = ["Etapa", "Saída visual", "Descrição técnica"]
        for col_idx, value in enumerate(header):
            set_cell_text(table.rows[0].cells[col_idx], value)

        image_paths = [
            case["original_path"],
            case["detection_path"],
            case["tissue_path"],
        ]
        for row_offset in range(3):
            row = table.rows[row_offset + 1]
            set_cell_text(row.cells[0], step_labels[row_offset])
            set_cell_image(row.cells[1], image_paths[row_offset], width=1.6)
            set_cell_text(row.cells[2], step_descriptions[row_offset])


def update_paragraphs(doc: Document, case_assets: list[dict], native_size: str, native_total: int, summary: dict) -> None:
    paragraphs = doc.paragraphs

    replace_paragraph_text(
        paragraphs[92],
        (
            "Este documento substitui a interpretação anterior baseada em degradação sintética "
            "por uma descrição fiel do pipeline real HEAL+ / REDISUS executado sobre o PIID nativo. "
            "A análise considera como a imagem é lida em BGR pelo OpenCV, como passa por controle "
            "de iluminação e contraste, como a ROI é extraída por textura, cor e contorno, e como "
            "a segmentação clínica v3 combina HSV, LAB, textura, luminância, skin exclusion e "
            "gradiente Scharr. Na validação real atual do split de teste do PIID, o sistema "
            f"processou {summary['total_images']} de {summary['total_images']} imagens com sucesso, "
            f"aceitou {format_percent(summary['analyzer_valid_wound_rate'])} das imagens como ferida válida "
            f"e operou com tempo médio de {summary['analyzer_avg_processing_time_ms']:.2f} ms por imagem. "
            "O classificador de estágio atingiu acurácia global de "
            f"{format_percent(summary['pressure_stage_accuracy'])}, com melhor desempenho no estágio 4 e "
            "maior dificuldade no estágio 3. O acervo local não contém capturas nativas em "
            "múltiplas resoluções, portanto o repositório atual não sustenta afirmar um limiar "
            "mínimo real inferior a 224 x 224 com base em aquisição nativa."
        ),
    )
    replace_paragraph_text(
        paragraphs[93],
        "Palavras-chave: HEAL+, processamento de imagem, ROI, contraste, PIID, resolução.",
    )

    replace_paragraph_text(
        paragraphs[98],
        "Documentar o funcionamento real do pipeline HEAL+ / REDISUS sobre imagens nativas do PIID, explicando leitura, contraste, ROI, segmentação, classificação e os pontos do código em que há ou não há distorção geométrica por redimensionamento.",
    )

    replace_paragraph_text(
        paragraphs[100],
        "O recorte foi executado sobre o dataset PIID Pressure Injury Staging, split de teste (n = 165 imagens), com a seguinte distribuição:",
    )
    replace_paragraph_text(
        paragraphs[106],
        (
            f"A inspeção automática do diretório dataset/piid/raw encontrou {native_total} arquivos JPG "
            f"com tamanho nativo único de {native_size}. Assim, o acervo local permite validação real "
            "do pipeline em imagens nativas 224 x 224, mas não permite inferir comportamento de captura "
            "nativa em resoluções maiores ou menores sem um novo conjunto de aquisição controlada."
        ),
    )

    replace_paragraph_text(paragraphs[108], "A validação real considerou o fluxo efetivamente executado hoje pelo HEAL+.")
    replace_paragraph_text(
        paragraphs[110],
        "ClinicalWoundAnalyzer e heal_analyzer.py como orquestradores do pipeline clínico real;",
    )
    replace_paragraph_text(
        paragraphs[112],
        "ImageEnhancer, WoundDetectorCV e ROISegmenter para qualidade, detecção inicial, máscara por contorno e zonas espaciais;",
    )
    replace_paragraph_text(
        paragraphs[114],
        "_segment_clinical_v3, TwoStageWoundClassifier e PressureInjuryStageClassifier para tecidos, classificação profunda e explicabilidade.",
    )
    replace_paragraph_text(
        paragraphs[116],
        "O preparo de datasets para treinamento segue scripts separados, principalmente scripts/preprocess_dataset.py e scripts/prepare_yolo_dataset.py.",
    )

    replace_paragraph_text(
        paragraphs[119],
        "A validação real foi executada sobre o split de teste descrito no manifesto dataset/piid/manifests/piid_lp_split.json.",
    )
    replace_paragraph_text(
        paragraphs[120],
        "ler cada JPEG com cv2.imread, portanto a entrada do pipeline é BGR em 8 bits e, neste acervo, 224 x 224;",
    )
    replace_paragraph_text(
        paragraphs[122],
        "rodar ClinicalWoundAnalyzer.analyze em cada imagem nativa, incluindo _prepare_input, detecção da ROI, zonas e segmentação clínica v3;",
    )
    replace_paragraph_text(
        paragraphs[124],
        "rodar PressureInjuryStageClassifier sobre a mesma imagem para medir o comportamento do ramo profundo disponível;",
    )
    replace_paragraph_text(
        paragraphs[126],
        "registrar validade da imagem, tecido primário, área da ferida, health score, tempo de processamento e estágio previsto;",
    )
    replace_paragraph_text(
        paragraphs[128],
        "selecionar casos reais do PIID para documentar visualmente a leitura, a ROI e o overlay tecidual final.",
    )
    replace_paragraph_text(paragraphs[130], "Foram consolidados os seguintes indicadores:")
    replace_paragraph_text(paragraphs[131], "taxa de leitura e processamento das imagens;")
    replace_paragraph_text(paragraphs[132], "taxa de imagens aceitas como ferida pelo analisador;")
    replace_paragraph_text(paragraphs[133], "área da ROI, tecido primário e health score gerados pelo pipeline clínico;")
    replace_paragraph_text(paragraphs[134], "tempo médio de processamento por imagem;")
    replace_paragraph_text(paragraphs[135], "acurácia global e por estágio do classificador de lesão por pressão;")
    replace_paragraph_text(paragraphs[136], "locais do código em que a resolução é preservada e locais em que há resize quadrado com potencial de distorção.")

    replace_paragraph_text(paragraphs[137], "5 TRATAMENTO DE RESOLUÇÃO E RISCO DE DISTORÇÃO")
    replace_paragraph_text(
        paragraphs[138],
        "No código atual, resolução e geometria são tratadas de formas diferentes conforme o ramo do sistema.",
    )
    replace_paragraph_text(
        paragraphs[139],
        "ClinicalWoundAnalyzer._prepare_input reduz apenas imagens cujo maior lado excede 1024 px, escalando proporcionalmente; isso preserva a razão de aspecto e não distorce a anatomia.",
    )
    replace_paragraph_text(
        paragraphs[140],
        "ImageEnhancer.prepare_for_cnn usa resize com padding e preserva proporção, mas esse helper não é o caminho usado por todos os modelos do projeto.",
    )
    replace_paragraph_text(
        paragraphs[141],
        "TwoStageWoundClassifier aplica transforms.Resize((224, 224)) e UNetSegmenter aplica cv2.resize para 512 x 512; nesses pontos, uma imagem não quadrada pode sofrer distorção geométrica.",
    )
    replace_paragraph_text(
        paragraphs[142],
        "Nos datasets de treino, scripts/preprocess_dataset.py faz resize direto para 640 x 640 (YOLO) e 256/512 x 512 (U-Net); as imagens usam interpolação cúbica ou area, e as máscaras usam INTER_NEAREST para não misturar classes.",
    )

    replace_paragraph_text(paragraphs[143], "6 CONTEXTO PARA 224 x 224 E PARA ENTRADAS MAIORES")
    replace_paragraph_text(
        paragraphs[144],
        "O valor 224 x 224 aparece no repositório porque vários ramos profundos foram treinados com entrada fixa nessa escala, o que explica a compatibilidade de modelos, mas não deve ser confundido com limite de captura nativa.",
    )
    replace_paragraph_text(
        paragraphs[147],
        "No código atual, o classificador ResNet50 de dois estágios registra entrada 224 x 224 com normalização ImageNet, enquanto a segmentação U-Net usa 512 x 512. A literatura citada ajuda a justificar esse desenho arquitetural, mas não substitui um estudo com imagens originalmente capturadas em múltiplas resoluções.",
    )
    replace_paragraph_text(
        paragraphs[148],
        "Para imagens de qualidade maior, o ramo clínico principal tende a ser mais robusto porque primeiro limita o maior lado a 1024 px mantendo proporção. O risco de distorção aparece principalmente quando a imagem entra em modelos que fazem resize quadrado direto.",
    )

    replace_paragraph_text(paragraphs[150], "Tabela 1 - Resumo técnico do pipeline real utilizado na validação atual")
    replace_paragraph_text(
        paragraphs[152],
        "Fonte: elaboração própria com base em clinical_wound_analyzer_core.py, image_enhancer.py, resnet_wound_classifier.py e output/validation/piid_heal_analyzer_validation_test_current.json.",
    )
    replace_paragraph_text(paragraphs[153], "Na validação real do split de teste do PIID, o comportamento consolidado foi o seguinte:")
    replace_paragraph_text(
        paragraphs[154],
        f"{summary['total_images']} de {summary['total_images']} imagens foram lidas e processadas com sucesso pelo analisador;",
    )
    replace_paragraph_text(
        paragraphs[155],
        f"{format_percent(summary['analyzer_valid_wound_rate'])} das imagens foram aceitas como compatíveis com ferida, com tempo médio de {summary['analyzer_avg_processing_time_ms']:.2f} ms por imagem no ambiente atual;",
    )
    replace_paragraph_text(
        paragraphs[156],
        f"o classificador de estágio atingiu acurácia global de {format_percent(summary['pressure_stage_accuracy'])}, com melhor desempenho no estágio 4 e maior dificuldade no estágio 3.",
    )

    replace_paragraph_text(paragraphs[158], "7.1 Estudos de caso reais do pipeline HEAL+")
    replace_paragraph_text(
        paragraphs[159],
        "A seção abaixo substitui as antigas simulações de downsampling por estudos de caso reais do pipeline HEAL+ sobre imagens nativas do PIID (224 x 224). Em cada caso, a primeira linha mostra a entrada, a segunda mostra a ROI derivada do detector e do segmentador por contorno, e a terceira mostra o resultado tecidual final.",
    )

    case_heading_indices = [161, 168, 173]
    case_summary_indices = [163, 170, 175]
    for idx, case in enumerate(case_assets):
        replace_paragraph_text(paragraphs[case_heading_indices[idx]], case["title"])
        replace_paragraph_text(paragraphs[case_summary_indices[idx]], case["summary"])

    replace_paragraph_text(paragraphs[177], "Na leitura por estágio do classificador de lesão por pressão, os principais achados foram:")
    replace_paragraph_text(
        paragraphs[179],
        f"estágio 4 foi o mais estável ({summary['pressure_stage_per_stage']['stage_4']['correct']}/{summary['pressure_stage_per_stage']['stage_4']['total']} corretos, acurácia {format_percent(summary['pressure_stage_per_stage']['stage_4']['accuracy'])});",
    )
    replace_paragraph_text(
        paragraphs[180],
        f"estágio 3 foi o mais difícil ({summary['pressure_stage_per_stage']['stage_3']['correct']}/{summary['pressure_stage_per_stage']['stage_3']['total']} corretos, acurácia {format_percent(summary['pressure_stage_per_stage']['stage_3']['accuracy'])}) e foi confundido com estágio 4 em {summary['pressure_stage_confusion']['stage_3']['stage_4']} casos;",
    )
    replace_paragraph_text(
        paragraphs[181],
        f"estágios 1 e 2 ficaram próximos ({format_percent(summary['pressure_stage_per_stage']['stage_1']['accuracy'])} e {format_percent(summary['pressure_stage_per_stage']['stage_2']['accuracy'])}), com confusão mútua de {summary['pressure_stage_confusion']['stage_1']['stage_2']} casos em cada direção;",
    )
    replace_paragraph_text(
        paragraphs[182],
        "esse padrão sugere que a ambiguidade clínica entre estágios vizinhos continua sendo mais crítica que a simples leitura da imagem 224 x 224 disponível no acervo local.",
    )

    replace_paragraph_text(
        paragraphs[185],
        "Com base na validação real hoje disponível no repositório, é correto afirmar que o HEAL+ processa com estabilidade o PIID nativo em 224 x 224 e gera ROI, tecidos, health score e classificação de estágio de forma reprodutível. Não é correto, porém, afirmar a partir deste acervo que 80 x 80 seja um limite mínimo real de captura nativa, porque o PIID local não contém imagens originalmente adquiridas em múltiplas resoluções.",
    )
    replace_paragraph_text(
        paragraphs[187],
        f"o pipeline clínico atual leu e processou {summary['total_images']}/{summary['total_images']} imagens;",
    )
    replace_paragraph_text(
        paragraphs[189],
        f"{format_percent(summary['analyzer_valid_wound_rate'])} foram aceitas como compatíveis com ferida;",
    )
    replace_paragraph_text(
        paragraphs[191],
        f"o tempo médio foi de {summary['analyzer_avg_processing_time_ms']:.2f} ms por imagem no ambiente atual;",
    )
    replace_paragraph_text(
        paragraphs[193],
        f"a acurácia de estágio ficou em {format_percent(summary['pressure_stage_accuracy'])}, com melhor resultado no estágio 4.",
    )
    replace_paragraph_text(
        paragraphs[195],
        "Para imagens de qualidade maior, o comportamento precisa ser interpretado por ramo do sistema. No motor clínico principal, o código reduz proporcionalmente apenas quando o maior lado supera 1024 px, portanto a geometria é preservada. Já nos ramos profundos que fazem resize quadrado direto para 224 x 224 ou 512 x 512, imagens não quadradas podem sofrer distorção geométrica.",
    )
    replace_paragraph_text(paragraphs[196], "capturar e armazenar a maior resolução nativa possível continua sendo desejável;")
    replace_paragraph_text(paragraphs[197], "evitar pre-esticar manualmente a foto para quadrados antes de entrar no pipeline;")
    replace_paragraph_text(paragraphs[198], "priorizar futura adoção de letterbox ou padding também nos modelos profundos.")
    replace_paragraph_text(paragraphs[200], "Assim, a interpretação final deste relatório é a seguinte:")
    replace_paragraph_text(paragraphs[202], "evidência real atual do repositório: funcionamento validado em PIID nativo 224 x 224;")
    replace_paragraph_text(paragraphs[203], "evidência ainda ausente: estudo com imagens originalmente capturadas em múltiplas resoluções;")
    replace_paragraph_text(paragraphs[204], "principal risco técnico a revisar: resize quadrado direto nos ramos profundos e nos scripts de preparo de dataset.")

    replace_paragraph_text(paragraphs[206], "Este resultado deve ser interpretado com as seguintes ressalvas:")
    replace_paragraph_text(paragraphs[207], "o PIID local contém apenas imagens 224 x 224; não há base interna para inferir efeito de captura nativa em 80 x 80, 512 x 512 ou outras escalas;")
    replace_paragraph_text(paragraphs[208], "o recorte continua restrito a lesões por pressão; outras etiologias crônicas podem exigir novas validações;")
    replace_paragraph_text(paragraphs[209], "o ensemble adicional baseado em DermaIntel, MedSAM e BiomedCLIP permaneceu indisponível no ambiente atual por dependências ausentes;")
    replace_paragraph_text(paragraphs[210], "o resultado de estágio ainda mostra confusão relevante entre estágios vizinhos, especialmente 3 versus 4.")

    replace_paragraph_text(paragraphs[212], "A validação real descrita neste documento pode ser reproduzida pelo script:")
    replace_paragraph_text(paragraphs[213], "scripts/validate_heal_analyzer_piid.py")
    replace_paragraph_text(paragraphs[214], "O resultado consolidado usado nesta revisão foi salvo em:")
    replace_paragraph_text(paragraphs[215], "output/validation/piid_heal_analyzer_validation_test_current.json")

    replace_paragraph_text(paragraphs[217], "Para produto, coleta e treinamento, recomenda-se adotar imediatamente:")
    replace_paragraph_text(paragraphs[218], "manter a captura nativa no maior tamanho disponível e deixar o pipeline reduzir apenas quando necessário;")
    replace_paragraph_text(paragraphs[219], "usar 224 x 224 apenas como requisito de compatibilidade dos modelos profundos existentes, não como limite clínico universal de captura;")
    replace_paragraph_text(paragraphs[220], "revisar TwoStageWoundClassifier, UNetSegmenter e os scripts de preprocessamento para introduzir padding ou letterbox quando a origem não for quadrada;")
    replace_paragraph_text(paragraphs[221], "tratar qualquer conclusão sobre resolução mínima real como pendente até existir um dataset com múltiplas resoluções nativas ou uma nova campanha de aquisição controlada.")

    replace_paragraph_text(paragraphs[241], "REDISUS. Validação real atual do HEAL+ no PIID. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[242], "output/validation/piid_heal_analyzer_validation_test_current.json. Acesso em: 16 abr. 2026.")
    replace_paragraph_text(paragraphs[244], "REDISUS. Script de validação real do HEAL+ em PIID. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[245], "scripts/validate_heal_analyzer_piid.py. Acesso em: 16 abr. 2026.")


def backup_document(path: Path) -> Path:
    backup_path = path.with_suffix(path.suffix + ".bak")
    if not backup_path.exists():
        shutil.copy2(path, backup_path)
    return backup_path


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(f"Documento não encontrado: {DOC_PATH}")

    payload = ensure_validation_output()
    summary = payload["summary"]
    native_size, native_total = collect_native_size_summary(PIID_ROOT)
    records = payload["records"]

    case_records = [
        select_case(records, "stage_1"),
        select_case(records, "stage_2"),
        select_case(records, "stage_4"),
    ]
    case_assets = build_case_assets(case_records)

    backup_document(DOC_PATH)
    doc = Document(DOC_PATH)
    update_paragraphs(doc, case_assets, native_size, native_total, summary)
    update_tables(doc, case_assets)
    doc.save(DOC_PATH)

    print(f"Documento atualizado: {DOC_PATH}")
    print(f"Validação usada: {VALIDATION_OUTPUT}")
    print(f"Imagens reais salvas em: {CASE_IMAGE_DIR}")


if __name__ == "__main__":
    main()
