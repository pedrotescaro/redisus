from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOC_PATH = ROOT / "docs" / "research" / "relatorio_redisus_resolucao_com_imagens.docx"
REPORT_PATH = ROOT / "docs" / "research" / "relatorio_medetec_pressao_resolucao_minima_operacional.docx"
VALIDATION_SCRIPT = ROOT / "scripts" / "validate_medetec_pressure_multiresolution.py"
VALIDATION_OUTPUT = ROOT / "output" / "validation" / "medetec_pressure_multiresolution_validation.json"
CASE_IMAGE_DIR = ROOT / "tmp_images" / "medetec_pressure_minres_cases"

if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.processing.clinical_wound_analyzer_core import ClinicalWoundAnalyzer


TEST_KEYS = ["native", "resized_512", "resized_320", "resized_224", "resized_160"]


def copy_run_format(source_run, target_run) -> None:
    if source_run is None:
        return
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.style = source_run.style
    if source_run.font is not None and target_run.font is not None:
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.color.rgb = source_run.font.color.rgb
        target_run.font.highlight_color = source_run.font.highlight_color
        target_run.font.superscript = source_run.font.superscript
        target_run.font.subscript = source_run.font.subscript
        target_run.font.all_caps = source_run.font.all_caps
        target_run.font.small_caps = source_run.font.small_caps


def replace_paragraph_text(paragraph, text: str) -> None:
    source_run = paragraph.runs[0] if paragraph.runs else None
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("pPr"):
            continue
        p.remove(child)
    new_run = paragraph.add_run(text)
    copy_run_format(source_run, new_run)


def clear_cell(cell) -> None:
    cell.text = ""


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    source_run = paragraph.runs[0] if paragraph.runs else None
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    new_run = paragraph.add_run(text)
    copy_run_format(source_run, new_run)


def set_cell_image(cell, image_path: Path, width: float = 2.2) -> None:
    clear_cell(cell)
    run = cell.paragraphs[0].add_run()
    run.add_picture(str(image_path), width=Inches(width))


def ensure_validation_output() -> dict:
    if not VALIDATION_OUTPUT.exists():
        cmd = [
            sys.executable,
            str(VALIDATION_SCRIPT),
            "--output",
            str(VALIDATION_OUTPUT.relative_to(ROOT)),
        ]
        subprocess.run(cmd, cwd=ROOT, check=True)
    return json.loads(VALIDATION_OUTPUT.read_text(encoding="utf-8"))


def format_percent(rate: float) -> str:
    return f"{rate * 100:.2f}% ({rate:.4f})"


def format_stage(stage_code: str) -> str:
    if not stage_code:
        return "desconhecido"
    return stage_code.replace("stage_", "estágio ")


def resolution_label(key: str) -> str:
    return "nativo" if key == "native" else key.replace("resized_", "") + " x " + key.replace("resized_", "")


def compute_operational_floor(payload: dict) -> tuple[int | None, dict[str, float]]:
    criteria = {
        "primary_tissue_agreement_rate": 0.90,
        "predicted_stage_agreement_rate": 0.90,
        "avg_abs_health_score_delta": 6.0,
        "avg_abs_wound_fraction_delta": 0.09,
    }
    floor = None
    for size in [512, 448, 320, 256, 224, 192, 160]:
        summary = payload["summary"]["vs_native"][f"resized_{size}"]
        if (
            summary["primary_tissue_agreement_rate"] >= criteria["primary_tissue_agreement_rate"]
            and summary["predicted_stage_agreement_rate"] >= criteria["predicted_stage_agreement_rate"]
            and summary["avg_abs_health_score_delta"] <= criteria["avg_abs_health_score_delta"]
            and summary["avg_abs_wound_fraction_delta"] <= criteria["avg_abs_wound_fraction_delta"]
        ):
            floor = size
    return floor, criteria


def select_cases(records: list[dict]) -> list[dict]:
    stable_all = []
    break_224 = []
    severe_160 = []

    for record in records:
        native = record["variants"]["native"]

        same_all = True
        for key in ["resized_512", "resized_448", "resized_320", "resized_256", "resized_224", "resized_192", "resized_160"]:
            variant = record["variants"][key]
            if (
                native.get("analyzer_primary_tissue") != variant.get("analyzer_primary_tissue")
                or native.get("predicted_stage") != variant.get("predicted_stage")
            ):
                same_all = False
                break
        if same_all:
            stable_all.append(record)

        stable_320 = (
            native.get("analyzer_primary_tissue") == record["variants"]["resized_320"].get("analyzer_primary_tissue")
            and native.get("predicted_stage") == record["variants"]["resized_320"].get("predicted_stage")
        )
        break_at_224 = (
            native.get("analyzer_primary_tissue") != record["variants"]["resized_224"].get("analyzer_primary_tissue")
            or native.get("predicted_stage") != record["variants"]["resized_224"].get("predicted_stage")
        )
        if stable_320 and break_at_224:
            break_224.append(record)

        severity = 0.0
        severity += abs(
            float(native.get("analyzer_health_score") or 0.0)
            - float(record["variants"]["resized_160"].get("analyzer_health_score") or 0.0)
        )
        severity += 100 * abs(
            float(native.get("analyzer_wound_fraction") or 0.0)
            - float(record["variants"]["resized_160"].get("analyzer_wound_fraction") or 0.0)
        )
        if (
            native.get("analyzer_primary_tissue") != record["variants"]["resized_160"].get("analyzer_primary_tissue")
            or native.get("predicted_stage") != record["variants"]["resized_160"].get("predicted_stage")
        ):
            severity += 10
        severe_160.append((severity, record))

    stable_case = max(
        stable_all,
        key=lambda record: (
            abs(float(record["variants"]["native"]["analyzer_health_score"]) - float(record["variants"]["resized_160"]["analyzer_health_score"])),
            abs(float(record["variants"]["native"]["analyzer_wound_fraction"]) - float(record["variants"]["resized_160"]["analyzer_wound_fraction"])),
        ),
    )
    threshold_case = max(
        break_224,
        key=lambda record: (
            abs(float(record["variants"]["native"]["analyzer_health_score"]) - float(record["variants"]["resized_224"]["analyzer_health_score"])),
            abs(float(record["variants"]["native"]["analyzer_wound_fraction"]) - float(record["variants"]["resized_224"]["analyzer_wound_fraction"])),
        ),
    )
    severe_case = max(severe_160, key=lambda item: item[0])[1]
    return [stable_case, threshold_case, severe_case]


def resize_to_height(image: np.ndarray, height: int) -> np.ndarray:
    h, w = image.shape[:2]
    new_width = max(1, int(round(w * (height / max(1, h)))))
    return cv2.resize(image, (new_width, height), interpolation=cv2.INTER_AREA)


def make_montage(label: str, original: np.ndarray, detection: np.ndarray, tissue: np.ndarray) -> np.ndarray:
    target_height = 160
    panels = [
        resize_to_height(original, target_height),
        resize_to_height(detection, target_height),
        resize_to_height(tissue, target_height),
    ]
    separator = np.full((target_height, 8, 3), 255, dtype=np.uint8)
    body = panels[0]
    for panel in panels[1:]:
        body = np.hstack([body, separator, panel])

    header = np.full((34, body.shape[1], 3), 255, dtype=np.uint8)
    cv2.putText(header, label, (10, 23), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (20, 20, 20), 2, cv2.LINE_AA)
    return np.vstack([header, body])


def build_case_assets(case_records: list[dict]) -> list[dict]:
    CASE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    analyzer = ClinicalWoundAnalyzer()
    assets: list[dict] = []
    titles = [
        "Caso de estudo 1: estabilidade dos rótulos com oscilação das métricas de acompanhamento",
        "Caso de estudo 2: estabilidade até 320 x 320 e divergência em 224 x 224",
        "Caso de estudo 3: degradação acentuada em 160 x 160",
    ]

    for index, record in enumerate(case_records, start=1):
        image_path = Path(str(record["path"]))
        image = cv2.imread(str(image_path))
        if image is None:
            raise FileNotFoundError(f"Não foi possível carregar {image_path}")

        montage_paths: dict[str, Path] = {}
        for key in TEST_KEYS:
            if key == "native":
                variant_image = image
            else:
                size = int(key.replace("resized_", ""))
                variant_image = cv2.resize(image, (size, size), interpolation=cv2.INTER_AREA)
            report = analyzer.analyze(variant_image)
            montage = make_montage(
                resolution_label(key),
                variant_image,
                report.detection_overlay,
                report.tissue_overlay,
            )
            out_path = CASE_IMAGE_DIR / f"case{index}_{key}.png"
            cv2.imwrite(str(out_path), montage)
            montage_paths[key] = out_path

        assets.append(
            {
                "record": record,
                "title": titles[index - 1],
                "montage_paths": montage_paths,
            }
        )
    return assets


def update_tables(doc: Document, case_assets: list[dict], payload: dict) -> None:
    dataset = payload["dataset"]
    resolutions = payload["summary"]["resolutions"]
    vs_native = payload["summary"]["vs_native"]

    table = doc.tables[0]
    rows = [
        ["Resolução", "Retenção média de pixels", "Acordo com o nativo", "Impacto nas métricas", "Leitura prática"],
        [
            "Nativo",
            "100.00%",
            "referência",
            f"health médio {resolutions['native']['avg_health_score']:.2f}; fração média da ROI {resolutions['native']['avg_wound_fraction'] * 100:.2f}%",
            f"{dataset['total_images']} imagens úteis entre {dataset['min_size']['width']} x {dataset['min_size']['height']} e {dataset['max_size']['width']} x {dataset['max_size']['height']}",
        ],
    ]
    interpretations = {
        512: "preservou de forma mais fiel os detalhes de interesse em relação ao nativo",
        448: "manteve elevada proximidade em relação ao resultado nativo",
        320: "permaneceu estável para uso operacional, com deriva moderada",
        256: "ingressou em faixa de instabilidade tecidual relevante",
        224: "não se mostrou equivalente ao nativo para fins de acompanhamento",
        192: "manteve validação de ferida, mas com oscilação excessiva para monitoramento seguro",
        160: "apresentou degradação acentuada; inadequado para seguimento clínico automatizado",
    }
    for size in [512, 448, 320, 256, 224, 192, 160]:
        key = f"resized_{size}"
        res = vs_native[key]
        rows.append(
            [
                f"{size} x {size}",
                f"{res['avg_pixel_retention_ratio'] * 100:.2f}%",
                f"tecido primário {format_percent(res['primary_tissue_agreement_rate'])}; estágio previsto {format_percent(res['predicted_stage_agreement_rate'])}",
                f"variação absoluta média do health score {res['avg_abs_health_score_delta']:.2f}; variação absoluta média da ROI {res['avg_abs_wound_fraction_delta'] * 100:.2f} pp",
                interpretations[size],
            ]
        )

    for row_idx, row_values in enumerate(rows):
        if row_idx >= len(table.rows):
            table.add_row()
        row = table.rows[row_idx]
        for col_idx, value in enumerate(row_values):
            set_cell_text(row.cells[col_idx], value)

    row_labels = ["nativo", "512 x 512", "320 x 320", "224 x 224", "160 x 160"]
    key_map = ["native", "resized_512", "resized_320", "resized_224", "resized_160"]

    for table_index, case in enumerate(case_assets, start=1):
        case_table = doc.tables[table_index]
        while len(case_table.rows) < 6:
            case_table.add_row()
        header = ["Variante", "Saída visual", "Descrição técnica"]
        for col_idx, value in enumerate(header):
            set_cell_text(case_table.rows[0].cells[col_idx], value)

        record = case["record"]
        for idx, key in enumerate(key_map):
            variant = record["variants"][key]
            row = case_table.rows[idx + 1]
            set_cell_text(row.cells[0], row_labels[idx])
            set_cell_image(row.cells[1], case["montage_paths"][key], width=2.15)
            if key == "native":
                description = (
                    f"Imagem original {record['native_width']} x {record['native_height']}. "
                    f"Tecido {variant['analyzer_primary_tissue']}; estágio {format_stage(variant['predicted_stage'])}; "
                    f"health score {variant['analyzer_health_score']:.2f}; ROI {variant['analyzer_wound_fraction'] * 100:.2f}%."
                )
            else:
                description = (
                    f"{resolution_label(key)}. "
                    f"Tecido {variant['analyzer_primary_tissue']}; estágio {format_stage(variant['predicted_stage'])}; "
                    f"health score {variant['analyzer_health_score']:.2f}; ROI {variant['analyzer_wound_fraction'] * 100:.2f}%."
                )
            set_cell_text(row.cells[2], description)


def update_paragraphs(doc: Document, case_assets: list[dict], payload: dict) -> None:
    paragraphs = doc.paragraphs
    dataset = payload["dataset"]
    resolutions = payload["summary"]["resolutions"]
    vs_native = payload["summary"]["vs_native"]
    floor, criteria = compute_operational_floor(payload)

    replace_paragraph_text(
        paragraphs[12],
        "Relatório Técnico: Avaliação da Resolução Mínima Operacional para Análise de Lesões por Pressão em Imagens do Subconjunto Medetec",
    )
    replace_paragraph_text(paragraphs[29], "17 de abril de 2026")
    replace_paragraph_text(paragraphs[63], "17 de abril de 2026")
    replace_paragraph_text(
        paragraphs[39],
        "Relatório Técnico: Avaliação da Resolução Mínima Operacional para Análise de Lesões por Pressão em Imagens do Subconjunto Medetec",
    )

    replace_paragraph_text(
        paragraphs[92],
        (
            "Este relatório examina em que medida a redução da resolução espacial compromete a análise automatizada e "
            "o acompanhamento de lesões por pressão. No subconjunto Medetec, foram comparadas 175 imagens nativas, "
            "com dimensões entre 559 x 347 e 560 x 560, e versões redimensionadas para 512, 448, 320, 256, 224, 192 e "
            "160 pixels quadrados. Em todas as resoluções, o HEAL+ manteve reconhecimento de ferida válida em 100% "
            "dos casos, porém a estabilidade clínica diminuiu progressivamente à medida que a resolução foi reduzida. "
            f"Em 512 x 512, observou-se concordância de {format_percent(vs_native['resized_512']['primary_tissue_agreement_rate'])} "
            f"para o tecido primário e de {format_percent(vs_native['resized_512']['predicted_stage_agreement_rate'])} "
            f"para o estágio previsto, com variação absoluta média de {vs_native['resized_512']['avg_abs_health_score_delta']:.2f} "
            "pontos no health score, indicando elevada preservação dos detalhes de interesse. "
            f"Em 224 x 224, observou-se concordância de {format_percent(vs_native['resized_224']['primary_tissue_agreement_rate'])} "
            f"para o tecido primário e de {format_percent(vs_native['resized_224']['predicted_stage_agreement_rate'])} "
            f"para o estágio previsto, com variação absoluta média de {vs_native['resized_224']['avg_abs_health_score_delta']:.2f} "
            "pontos no health score. Considerando-se um critério conservador de estabilidade operacional, definido por "
            f"concordância do tecido primário igual ou superior a {criteria['primary_tissue_agreement_rate']:.2f}, "
            f"concordância do estágio previsto igual ou superior a {criteria['predicted_stage_agreement_rate']:.2f}, "
            f"variação absoluta média do health score igual ou inferior a {criteria['avg_abs_health_score_delta']:.1f} "
            f"e variação absoluta média da ROI igual ou inferior a {criteria['avg_abs_wound_fraction_delta'] * 100:.1f} "
            f"pontos percentuais, verificou-se que a menor resolução ainda aceitável foi {floor} x {floor}. "
            "A partir de 256 x 256, observou-se instabilidade na leitura tecidual, e 224 x 224 não se mostrou "
            "equivalente ao nativo para fins de seguimento clínico."
        ),
    )
    replace_paragraph_text(
        paragraphs[93],
        "Palavras-chave: Medetec, lesão por pressão, resolução mínima, acompanhamento, escore de saúde, HEAL+.",
    )

    replace_paragraph_text(paragraphs[97], "1 OBJETIVO")
    replace_paragraph_text(
        paragraphs[98],
        "Determinar, no subconjunto Medetec de lesão por pressão, se a redução de imagens de maior resolução compromete a análise realizada pela IA, em que faixa esse impacto se torna clinicamente relevante e qual resolução mínima operacional pode ser recomendada para o pipeline atual.",
    )

    replace_paragraph_text(paragraphs[99], "2 RECORTE EXPERIMENTAL")
    replace_paragraph_text(
        paragraphs[100],
        "A base analisada foi composta exclusivamente pelas pastas Medetec pressure_ulcers_1 e pressure_ulcers_2, com deduplicação de cópias espelhadas e exclusão das introslides.",
    )
    replace_paragraph_text(paragraphs[101], "pressure_ulcers_1: 101 imagens úteis;")
    replace_paragraph_text(paragraphs[102], "pressure_ulcers_2: 74 imagens úteis;")
    replace_paragraph_text(paragraphs[103], "pressure_ulcers_set_1_of_2 e pressure_ulcers_set_2_of_2 foram descartadas por serem duplicatas hash-idênticas;")
    replace_paragraph_text(paragraphs[104], f"total final: {dataset['total_images']} imagens únicas, com {dataset['unique_sizes']} resoluções nativas entre {dataset['min_size']['width']} x {dataset['min_size']['height']} e {dataset['max_size']['width']} x {dataset['max_size']['height']}.")
    replace_paragraph_text(
        paragraphs[106],
        "Esse desenho permite avaliar a mesma ferida em várias resoluções artificiais derivadas de uma captura nativa maior, o que é suficiente para estimar um piso operacional do pipeline, embora não substitua um estudo prospectivo de aquisição clínica em múltiplas resoluções nativas.",
    )

    replace_paragraph_text(paragraphs[107], "3 DEFINIÇÃO DE PERDA DE QUALIDADE NO CONTEXTO DO ESTUDO")
    replace_paragraph_text(
        paragraphs[108],
        "No contexto deste estudo, perda de qualidade não corresponde apenas à redução do tamanho do arquivo. Refere-se à diminuição da amostragem espacial da lesão, com suavização de textura, atenuação de bordas e compressão geométrica para um formato quadrado fixo quando a imagem original não é quadrada.",
    )
    replace_paragraph_text(paragraphs[110], "à medida que a resolução diminui, a imagem preserva a presença global da lesão, porém perde microdetalhes de esfacelo, borda, brilho e textura;")
    replace_paragraph_text(paragraphs[112], "essa perda pode não impedir a validação binária de ferida, mas altera o tecido primário, a área relativa da ROI, o health score e o estágio previsto;")
    replace_paragraph_text(paragraphs[115], "desse modo, o impacto clínico mais relevante tende a ocorrer no acompanhamento longitudinal e na interpretação tecidual, e não apenas na detecção da presença de ferida.")
    replace_paragraph_text(
        paragraphs[117],
        "Assim, o objetivo experimental consiste em identificar a partir de que ponto essa perda passa a afetar de modo material a saída produzida pela IA.",
    )

    replace_paragraph_text(paragraphs[119], "4 METODOLOGIA")
    replace_paragraph_text(paragraphs[120], "Cada imagem foi analisada no nativo e em sete redimensionamentos quadrados: 512, 448, 320, 256, 224, 192 e 160.")
    replace_paragraph_text(paragraphs[121], "ler a imagem nativa com cv2.imread e executá-la no ClinicalWoundAnalyzer sem redução interna adicional;")
    replace_paragraph_text(paragraphs[122], "gerar variantes quadradas por cv2.resize com INTER_AREA, preservando conteúdo global, mas não a razão de aspecto;")
    replace_paragraph_text(paragraphs[123], "medir, em cada variante, ferida válida, tecido primário, fração da ROI, health score, tempo de processamento e estágio previsto;")
    replace_paragraph_text(paragraphs[124], "comparar cada resolução diretamente contra o nativo para medir concordância e deltas clínicos;")
    replace_paragraph_text(paragraphs[125], "definir uma resolução mínima operacional pela menor escala que ainda preserva concordância e estabilidade dentro de limites conservadores.")
    replace_paragraph_text(paragraphs[127], "Foram usados os seguintes critérios de análise:")
    replace_paragraph_text(paragraphs[128], "concordância do tecido primário com a imagem nativa;")
    replace_paragraph_text(paragraphs[129], "concordância do estágio previsto com a imagem nativa;")
    replace_paragraph_text(paragraphs[130], "delta absoluto médio de health score;")
    replace_paragraph_text(paragraphs[131], "delta absoluto médio da fração da ROI;")
    replace_paragraph_text(paragraphs[132], "retenção média de pixels em relação ao nativo;")
    replace_paragraph_text(paragraphs[133], "estudos de caso que mostram estabilidade, quebra em 224 e degradação severa em 160.")

    replace_paragraph_text(paragraphs[134], "5 RESULTADOS CONSOLIDADOS")
    replace_paragraph_text(
        paragraphs[135],
        "A Tabela 1 sintetiza a curva de estabilidade da IA à medida que a resolução é reduzida.",
    )
    replace_paragraph_text(
        paragraphs[136],
        f"512 x 512 reteve em média {vs_native['resized_512']['avg_pixel_retention_ratio'] * 100:.2f}% dos pixels do nativo e preservou de forma mais fiel os detalhes de interesse observados na imagem original;",
    )
    replace_paragraph_text(
        paragraphs[137],
        f"448 x 448 reteve {vs_native['resized_448']['avg_pixel_retention_ratio'] * 100:.2f}% dos pixels do nativo e manteve elevada proximidade em relação à leitura original;",
    )
    replace_paragraph_text(
        paragraphs[138],
        f"320 x 320 reteve {vs_native['resized_320']['avg_pixel_retention_ratio'] * 100:.2f}% dos pixels e ainda manteve concordância de {format_percent(vs_native['resized_320']['primary_tissue_agreement_rate'])} para o tecido e de {format_percent(vs_native['resized_320']['predicted_stage_agreement_rate'])} para o estágio;",
    )
    replace_paragraph_text(
        paragraphs[139],
        f"256 x 256 e resoluções inferiores deixaram de apresentar estabilidade para acompanhamento tecidual, com queda para {format_percent(vs_native['resized_256']['primary_tissue_agreement_rate'])} em 256 e {format_percent(vs_native['resized_224']['primary_tissue_agreement_rate'])} em 224; em 160 x 160, o tecido primário mudou em {vs_native['resized_160']['primary_tissue_changed_count']} de {vs_native['resized_160']['total_pairs']} casos e a variação absoluta média do health score atingiu {vs_native['resized_160']['avg_abs_health_score_delta']:.2f}.",
    )

    replace_paragraph_text(paragraphs[140], "6 DISCUSSÃO DOS RESULTADOS")
    replace_paragraph_text(
        paragraphs[141],
        "Os resultados indicam que a perda de qualidade produz ao menos três efeitos distintos. Em primeiro lugar, a validação binária de ferida permaneceu estável em 100% dos pares, inclusive em 160 x 160, enquanto 512 x 512 e 448 x 448 mantiveram comportamento muito próximo ao nativo.",
    )
    replace_paragraph_text(
        paragraphs[142],
        f"Em segundo lugar, a interpretação tecidual passou a se degradar antes da validação binária: a concordância do tecido foi de {format_percent(vs_native['resized_512']['primary_tissue_agreement_rate'])} em 512, {format_percent(vs_native['resized_448']['primary_tissue_agreement_rate'])} em 448, {format_percent(vs_native['resized_320']['primary_tissue_agreement_rate'])} em 320, {format_percent(vs_native['resized_256']['primary_tissue_agreement_rate'])} em 256 e {format_percent(vs_native['resized_224']['primary_tissue_agreement_rate'])} em 224.",
    )
    replace_paragraph_text(
        paragraphs[143],
        f"Em terceiro lugar, o acompanhamento quantitativo também apresentou deriva: a variação absoluta média do health score foi de {vs_native['resized_512']['avg_abs_health_score_delta']:.2f} em 512, {vs_native['resized_448']['avg_abs_health_score_delta']:.2f} em 448, {vs_native['resized_320']['avg_abs_health_score_delta']:.2f} em 320, {vs_native['resized_256']['avg_abs_health_score_delta']:.2f} em 256, {vs_native['resized_224']['avg_abs_health_score_delta']:.2f} em 224 e {vs_native['resized_160']['avg_abs_health_score_delta']:.2f} em 160.",
    )

    replace_paragraph_text(paragraphs[145], "7 TABELA DE ESTABILIDADE")
    replace_paragraph_text(paragraphs[146], "Tabela 1 - Estabilidade do HEAL+ na Medetec conforme a resolução é reduzida")
    replace_paragraph_text(
        paragraphs[148],
        "Fonte: elaboração própria com base em dataset/medetec/metadata.json e output/validation/medetec_pressure_multiresolution_validation.json.",
    )
    replace_paragraph_text(
        paragraphs[149],
        "Os resultados indicam que imagens de menor resolução ainda permitem a identificação da presença de lesão, mas deixam de manter equivalência na quantificação tecidual, na extensão relativa da ROI e na leitura global de saúde da lesão.",
    )
    replace_paragraph_text(
        paragraphs[150],
        "512 x 512 correspondeu à faixa de maior proximidade em relação ao resultado obtido na imagem nativa, com melhor preservação dos detalhes de interesse;",
    )
    replace_paragraph_text(
        paragraphs[151],
        "448 x 448 também permaneceu muito próximo do nativo, embora com leve aumento de deriva quando comparado a 512 x 512;",
    )
    replace_paragraph_text(
        paragraphs[152],
        "320 x 320 correspondeu à menor resolução que ainda preservou estabilidade operacional compatível com os critérios adotados, ao passo que 224 x 224 não se mostrou equivalente ao nativo para fins de acompanhamento.",
    )

    replace_paragraph_text(paragraphs[153], "7.1 Estudos de caso reais da degradação por resolução")
    replace_paragraph_text(
        paragraphs[154],
        "As tabelas a seguir apresentam três trajetórias representativas. Em cada caso, são exibidas a imagem nativa e as versões em 512 x 512, 320 x 320, 224 x 224 e 160 x 160, permitindo observar a preservação ou a perda progressiva dos detalhes de interesse.",
    )
    case_heading_indices = [156, 166, 171]
    case_summary_indices = [158, 168, 173]
    summaries = []
    for case in case_assets:
        record = case["record"]
        native = record["variants"]["native"]
        v512 = record["variants"]["resized_512"]
        v320 = record["variants"]["resized_320"]
        v224 = record["variants"]["resized_224"]
        v160 = record["variants"]["resized_160"]
        summaries.append(
            (
                f"Arquivo: {record['filename']}. Nativo {record['native_width']} x {record['native_height']}. "
                f"Nativo: {native['analyzer_primary_tissue']} / {format_stage(native['predicted_stage'])} / health {native['analyzer_health_score']:.2f}. "
                f"512: {v512['analyzer_primary_tissue']} / {format_stage(v512['predicted_stage'])} / health {v512['analyzer_health_score']:.2f}. "
                f"320: {v320['analyzer_primary_tissue']} / {format_stage(v320['predicted_stage'])} / health {v320['analyzer_health_score']:.2f}. "
                f"224: {v224['analyzer_primary_tissue']} / {format_stage(v224['predicted_stage'])} / health {v224['analyzer_health_score']:.2f}. "
                f"160: {v160['analyzer_primary_tissue']} / {format_stage(v160['predicted_stage'])} / health {v160['analyzer_health_score']:.2f}."
            )
        )
    for idx, case in enumerate(case_assets):
        replace_paragraph_text(paragraphs[case_heading_indices[idx]], case["title"])
        replace_paragraph_text(paragraphs[case_summary_indices[idx]], summaries[idx])

    replace_paragraph_text(paragraphs[174], "8 DETERMINAÇÃO DA RESOLUÇÃO MÍNIMA OPERACIONAL")
    replace_paragraph_text(
        paragraphs[175],
        f"Com base no critério conservador explicitado na metodologia, a menor resolução testada que permaneceu dentro dos limites de estabilidade estabelecidos foi {floor} x {floor}.",
    )
    replace_paragraph_text(
        paragraphs[177],
        "512 x 512 apresentou a maior aderência ao resultado obtido na imagem nativa, preservando com maior fidelidade os detalhes de interesse, e 448 x 448 também permaneceu em faixa de elevada proximidade;",
    )
    replace_paragraph_text(
        paragraphs[178],
        "320 x 320 foi a menor resolução a manter concordância do tecido primário igual ou superior a 0,90, concordância de estágio igual ou superior a 0,90, variação absoluta média do health score igual ou inferior a 6,0 e variação absoluta média da ROI igual ou inferior a 9,0 pontos percentuais;",
    )
    replace_paragraph_text(
        paragraphs[179],
        "256 x 256 deixou de atender ao critério estabelecido em razão da instabilidade do tecido primário e do aumento relevante do health score;",
    )
    replace_paragraph_text(
        paragraphs[180],
        "224 x 224 não se mostrou adequado como piso para seguimento clínico neste pipeline, pois apresentou deriva relevante em comparação com a imagem nativa de maior resolução.",
    )

    replace_paragraph_text(paragraphs[182], "9 CONCLUSÃO E RECOMENDAÇÃO")
    replace_paragraph_text(
        paragraphs[183],
        "Os resultados obtidos demonstram que a redução da resolução implica perda de informação relevante para a análise automatizada. Esse efeito não se manifesta inicialmente na simples validação binária da presença de ferida, mas sobretudo na interpretação tecidual, na fração da ROI e no health score. Desse modo, verificou-se que a perda de detalhe pode comprometer o acompanhamento das feridas mesmo quando a detecção binária permanece funcional.",
    )
    replace_paragraph_text(paragraphs[185], "para o experimento realizado com o subconjunto Medetec, a resolução mínima operacional recomendada foi 320 x 320, enquanto 512 x 512 correspondeu à condição de maior preservação dos detalhes de interesse;")
    replace_paragraph_text(paragraphs[187], "256 x 256 já ingressa em faixa de instabilidade clínica, e 224 x 224 não se mostrou equivalente à imagem de maior resolução;")
    replace_paragraph_text(paragraphs[189], "para fins de acompanhamento longitudinal, health score e leitura tecidual, recomenda-se preservar a resolução nativa sempre que possível;")
    replace_paragraph_text(paragraphs[191], "quando houver exigência de entrada fixa, recomenda-se priorizar estratégias de adequação geométrica, como preenchimento de borda (padding) ou enquadramento com preservação de proporção (letterbox), em vez de achatamento quadrado direto.")
    replace_paragraph_text(
        paragraphs[193],
        "Em síntese, a redução da resolução não inviabiliza a detecção da lesão, mas pode alterar de modo relevante a interpretação produzida pelo sistema. Tal comportamento justifica a definição de um piso operacional de resolução.",
    )
    replace_paragraph_text(paragraphs[194], "recomenda-se capturar e armazenar a maior resolução nativa disponível como melhor prática, podendo-se utilizar 512 x 512 como alternativa de alta preservação quando houver necessidade de padronização superior a 320 x 320;")
    replace_paragraph_text(paragraphs[195], f"recomenda-se adotar {floor} x {floor} como piso operacional do pipeline atual com base nos dados desta validação;")
    replace_paragraph_text(paragraphs[196], "recomenda-se evitar 224 x 224 como base para acompanhamento, salvo em situações sem alternativa e com a devida ressalva metodológica.")

    replace_paragraph_text(paragraphs[198], "A interpretação final deste relatório pode ser sintetizada nos seguintes pontos:")
    replace_paragraph_text(paragraphs[200], "evidência observada: a redução da imagem altera a leitura da IA antes de comprometer a detecção da ferida, embora 512 x 512 tenha mantido a maior proximidade com o resultado nativo;")
    replace_paragraph_text(paragraphs[201], f"recomendação operacional: {floor} x {floor} corresponde ao menor ponto testado com estabilidade considerada aceitável no pipeline;")
    replace_paragraph_text(paragraphs[202], "principal alerta metodológico: 224 x 224 apresentou deriva relevante e, portanto, não deve ser considerado equivalente à imagem nativa de maior resolução.")

    replace_paragraph_text(paragraphs[203], "10 LIMITAÇÕES")
    replace_paragraph_text(paragraphs[204], "Este resultado deve ser interpretado com as seguintes ressalvas:")
    replace_paragraph_text(paragraphs[205], "a Medetec não fornece verdade de referência (ground truth) de estágio nem máscara manual, de modo que a análise mede estabilidade relativa ao nativo e não acurácia clínica absoluta;")
    replace_paragraph_text(paragraphs[206], "o piso de 320 x 320 é operacional para este pipeline e para este experimento, não devendo ser interpretado como parâmetro universal para toda ferida ou toda IA;")
    replace_paragraph_text(paragraphs[207], "o estudo utiliza resize quadrado direto; estratégias com preenchimento de borda (padding) podem preservar mais informação e merecem avaliação específica;")
    replace_paragraph_text(paragraphs[208], "o ensemble adicional com DermaIntel, MedSAM e BiomedCLIP permaneceu indisponível no ambiente atual.")

    replace_paragraph_text(paragraphs[209], "11 REPRODUTIBILIDADE")
    replace_paragraph_text(paragraphs[210], "A validação multirresolução descrita neste documento pode ser reproduzida pelo script:")
    replace_paragraph_text(paragraphs[211], "scripts/validate_medetec_pressure_multiresolution.py")
    replace_paragraph_text(paragraphs[212], "O resultado consolidado usado nesta revisão foi salvo em:")
    replace_paragraph_text(paragraphs[213], "output/validation/medetec_pressure_multiresolution_validation.json")

    replace_paragraph_text(paragraphs[214], "12 ENCAMINHAMENTO PRÁTICO")
    replace_paragraph_text(paragraphs[215], "Para o sistema e para a coleta clínica, recomenda-se adotar imediatamente:")
    replace_paragraph_text(paragraphs[216], "preservar a imagem nativa maior como padrão de captura e armazenamento;")
    replace_paragraph_text(paragraphs[217], f"usar {floor} x {floor} como menor resolução operacional recomendada quando houver necessidade de padronização e considerar 512 x 512 quando se desejar maior preservação dos detalhes de interesse;")
    replace_paragraph_text(paragraphs[218], "evitar 224 x 224 como base para acompanhamento de evolução da ferida;")
    replace_paragraph_text(paragraphs[219], "priorizar, em desenvolvimentos futuros, a substituição do resize quadrado direto por enquadramento com preservação de proporção (letterbox) ou preenchimento de borda (padding).")

    replace_paragraph_text(paragraphs[228], "src/processing/tissue_analyzer.py. Acesso em: 17 abr. 2026.")
    replace_paragraph_text(paragraphs[231], "src/processing/wound_detector_cv.py. Acesso em: 17 abr. 2026.")
    replace_paragraph_text(paragraphs[234], "src/processing/roi_segmentation.py. Acesso em: 17 abr. 2026.")

    replace_paragraph_text(paragraphs[236], "REDISUS. Medetec: metadados do subconjunto de lesão por pressão. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[237], "dataset/medetec/metadata.json. Acesso em: 17 abr. 2026.")
    replace_paragraph_text(paragraphs[239], "REDISUS. Validação multirresolução da Medetec em lesão por pressão. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[240], "output/validation/medetec_pressure_multiresolution_validation.json. Acesso em: 17 abr. 2026.")
    replace_paragraph_text(paragraphs[242], "REDISUS. Script de validação multirresolução da Medetec em lesão por pressão. [S. l.], 2026. Arquivo local:")
    replace_paragraph_text(paragraphs[243], "scripts/validate_medetec_pressure_multiresolution.py. Acesso em: 17 abr. 2026.")


def create_report() -> Path:
    if not TEMPLATE_DOC_PATH.exists():
        raise FileNotFoundError(f"Modelo de relatório não encontrado: {TEMPLATE_DOC_PATH}")

    payload = ensure_validation_output()
    case_records = select_cases(payload["records"])
    case_assets = build_case_assets(case_records)

    target_path = REPORT_PATH
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        shutil.copy2(TEMPLATE_DOC_PATH, target_path)
    except PermissionError:
        target_path = REPORT_PATH.with_stem(REPORT_PATH.stem + "_atualizado")
        shutil.copy2(TEMPLATE_DOC_PATH, target_path)

    doc = Document(target_path)
    update_paragraphs(doc, case_assets, payload)
    update_tables(doc, case_assets, payload)
    doc.save(target_path)
    return target_path


def main() -> int:
    output_path = create_report()
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
