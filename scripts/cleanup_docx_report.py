import docx
from docx.oxml import OxmlElement

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def delete_table(table):
    tbl = table._element
    tbl.getparent().remove(tbl)
    table._tbl = table._element = None

def delete_generated_sections(doc):
    """
    Remove todos os parágrafos e tabelas a partir do primeiro cabeçalho
    'Exemplos Visuais', que foi acidentalmente inserido múltiplas vezes no final.
    E para por ali (pra não deletar Referências caso houvesse antes, 
    mas no nosso doc as Referências estão e nós anexamos só no final)
    """
    elements_to_delete = []
    found_exemplos = False
    
    # Intera sobre a estrutura body do documento (XML block real)
    for element in doc._body._body:
        if element.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(element, doc)
            if para.text == 'Exemplos Visuais' or para.text == 'Exemplos e Simulações Visuais':
                found_exemplos = True
            
            if found_exemplos and "Referências" not in para.text:
                elements_to_delete.append(element)
            elif found_exemplos and "Referências" in para.text:
                # Se acharmos referências, interrompemos
                found_exemplos = False
                
        elif element.tag.endswith('tbl'):
            if found_exemplos:
                elements_to_delete.append(element)

    # Deletando do doc inteiro
    for el in elements_to_delete:
        el.getparent().remove(el)
        
    return doc

def reinsert_before_referencias(doc, real_img):
    import os
    from docx.shared import Inches
    from update_docx_report import process_image

    img_orig = process_image(real_img, 'tmp_images/orig.png', 224) 
    img_224 = process_image(real_img, 'tmp_images/roi_224.png', 224)
    img_80 = process_image(real_img, 'tmp_images/roi_80.png', 80)
    img_32 = process_image(real_img, 'tmp_images/roi_32.png', 32)
    
    # Encontrar as referencias (o ultimo titulo normal)
    ref_para = None
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Referências" in p.text:
            ref_para = p
            break
            
    if ref_para is None:
        # Se não achou Referências, insere no fim mesmo (seguro pra gerar novo)
        ref_para = doc.paragraphs[-1]
    
    # Inserir titulo
    new_p = ref_para.insert_paragraph_before("Exemplos e Simulações Visuais", style='Heading 1')
    
    # Inserir descricao
    desc = ref_para.insert_paragraph_before(
        "A fim de materializar qualitativamente o impacto discutido nos dados paramétricos, a Figura apresenta uma simulação da degradação espacial (simulando a captura por diferentes capacidades óticas, ou recortes profundos - downsampling - e posterior uniformização da entrada).\n\n"
        "O processo mostra o limite de reconhecimento tecidual e estabilidade da segmentação por ROI:\n"
        "  - Em 224x224 (Padrão e Limite Operacional): Texturas complexas finas, iluminação e delimitações irregulares dos tecidos esfacelados se mantêm visíveis.\n"
        "  - Em 80x80 (Limiar Estável - Mínimo Pipeline Clássico): Começa a ocorrer o arredondamento ou pixelização das bordas, todavia o modelo do sistema (WoundDetectorCV e TissueAnalyzerCV) não perde coesão central nem classifica excesso de espaço viável como falso neutro.\n"
        "  - Em 32x32 (Degradação Crítica): A diluição severa de pixels mescla o leito da ferida ao tom da pele externa e distorce gravemente focos de descoloração e brilho. A confiança dos contornos teciduais (Dice) despenca.\n\n"
        "Nota de Uso: Devido à licença estritamente voltada para fins pedagógicos e de validação em pesquisa acadêmica do conjunto PIID (Pressure Injury Image Dataset), estas imagens encontram-se restritas à composição metodológica deste relatório e vedadas a usos comerciais."
    )

    # Inserindo Tabela de imagens antes das Referências
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Imagem Original Reamostrada'
    hdr_cells[1].text = 'ROI'

    row_cells = table.add_row().cells
    p = row_cells[0].paragraphs[0]
    p.add_run('224x224 (Referência Inicial)\n').add_picture(img_orig, width=Inches(2.5))
    row_cells[1].paragraphs[0].add_run('ROI 224x224\n').add_picture(img_224, width=Inches(2.5))

    row_cells = table.add_row().cells
    p = row_cells[0].paragraphs[0]
    p.add_run('80x80 (Limiar Estável)\n').add_picture(img_orig, width=Inches(2.5))
    row_cells[1].paragraphs[0].add_run('ROI 80x80\n').add_picture(img_80, width=Inches(2.5))

    row_cells = table.add_row().cells
    p = row_cells[0].paragraphs[0]
    p.add_run('32x32 (Distorção Completa)\n').add_picture(img_orig, width=Inches(2.5))
    row_cells[1].paragraphs[0].add_run('ROI 32x32\n').add_picture(img_32, width=Inches(2.5))
    
    # Movemos XML element para antes das Referencias
    ref_para._element.addprevious(table._element)


if __name__ == "__main__":
    from docx import Document
    import glob
    doc_path = r'docs\research\relatorio_redisus_resolucao.docx'
    
    real_img = glob.glob(r'dataset\piid\**\*.jpg', recursive=True)[0]
    
    doc = Document(doc_path)
    doc = delete_generated_sections(doc)
    reinsert_before_referencias(doc, real_img)
    doc.save(doc_path)
    print("Documento DOCX limpo e reformatado corretamente.")
